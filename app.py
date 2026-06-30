# CRITICAL: Monkey patch MUST be first, before any other imports
from gevent import monkey
monkey.patch_all()

from flask import Flask, render_template, request, Response, jsonify, session  # type: ignore
import json
import os
import time
import gevent
from gevent import queue as gevent_queue
import psycopg2
import re
from search import SemanticSearch
from conversation_manager import conversation_manager
from freud_engine import get_engine as get_freud_engine
from jung_engine import get_engine as get_jung_engine
from kuczynski_engine import get_engine as get_kuczynski_engine
from hume_engine import get_engine as get_hume_engine
from nietzsche_engine import get_engine as get_nietzsche_engine
from bergler_engine import get_engine as get_bergler_engine

try:
    import anthropic  # type: ignore
    from anthropic import Anthropic  # type: ignore
except ImportError:
    print("Anthropic library not found. Installing...")
    Anthropic = None

try:
    from openai import OpenAI  # type: ignore
except ImportError:
    print("OpenAI library not found. Installing...")
    OpenAI = None

try:
    import PyPDF2  # type: ignore
except ImportError:
    print("PyPDF2 not found")
    PyPDF2 = None

try:
    import docx  # type: ignore
except ImportError:
    print("python-docx not found")
    docx = None

import re


def detect_explicit_requirements(question, default_length, default_quotes):
    """
    Detect when user's question explicitly requires more content than settings allow.
    Override settings when question asks for specific numbers of quotes, examples, etc.
    """
    question_lower = question.lower()

    override_length = default_length
    override_quotes = default_quotes

    quote_patterns = [
        r'(\d+)\s*(?:quotation|quote|citation|passage|excerpt)s?',
        r'(?:give|list|provide|show)\s*(?:me\s*)?(\d+)\s*(?:quotation|quote|example|citation|passage|excerpt)s?',
        r'(\d+)\s*(?:example|instance|case|illustration)s?\s*(?:from|of)',
    ]

    for pattern in quote_patterns:
        match = re.search(pattern, question_lower)
        if match:
            requested_quotes = int(match.group(1))
            if requested_quotes > default_quotes:
                override_quotes = min(requested_quotes, 100)
                words_per_quote = 50
                override_length = max(override_length, requested_quotes * words_per_quote)
                break

    length_patterns = [
        r'(\d+)\s*(?:word|sentence|paragraph)s?\s*(?:each|per|for each)',
        r'(?:at least|minimum)\s*(\d+)\s*words?',
        r'(\d+)\s*(?:sentence|paragraph)s?\s*(?:explanation|analysis|discussion)',
    ]

    for pattern in length_patterns:
        match = re.search(pattern, question_lower)
        if match:
            num = int(match.group(1))
            if 'sentence' in question_lower:
                estimated_words = num * 20 * max(override_quotes, 10)
            elif 'paragraph' in question_lower:
                estimated_words = num * 100 * max(override_quotes, 5)
            else:
                estimated_words = num
            override_length = max(override_length, estimated_words)
            break

    multi_part_indicators = [
        r'(?:answer|respond to)\s*(?:each|all)\s*(?:of\s*)?(?:the\s*)?(?:following|these)\s*(\d+)?',
        r'(\d+)\s*(?:question|query|item)s?',
        r'(?:for each|each of|all of)\s*(?:the\s*)?(\d+)?',
    ]

    question_count = question_lower.count('?')
    if question_count >= 2:
        per_question_words = 200
        sentence_match = re.search(r'(\d+)\s*(?:sentence|line)s?\s*(?:answer|response|reply)', question_lower)
        if sentence_match:
            per_question_words = max(per_question_words, int(sentence_match.group(1)) * 25)
        override_length = max(override_length, question_count * per_question_words)
        override_quotes = max(override_quotes, question_count * 2)

    for pattern in multi_part_indicators:
        match = re.search(pattern, question_lower)
        if match:
            if match.group(1):
                num_parts = int(match.group(1))
            else:
                num_parts = max(5, question_count)
            override_length = max(override_length, num_parts * 200)
            override_quotes = max(override_quotes, num_parts)
            break

    complex_question_indicators = [
        (r'\b(?:compare|contrast|differentiate|distinguish)\b.*\b(?:and|with|from|between)\b', 500),
        (r'\b(?:explain|describe|elaborate|discuss)\b.*\b(?:in detail|thoroughly|fully|comprehensively)\b', 500),
        (r'\b(?:what are|list|enumerate|name)\b.*\b(?:all|every|each|the main|the key|the primary)\b', 400),
        (r'\b(?:how does|how do|how is|how are)\b.*\b(?:relate|connect|influence|affect|impact)\b', 400),
        (r'\b(?:analyze|examine|evaluate|assess|critique)\b', 400),
        (r'\b(?:what is the relationship|what are the differences|what are the similarities)\b', 450),
        (r'\b(?:summarize|outline|overview)\b.*\b(?:theory|argument|position|view|philosophy)\b', 400),
        (r'\b(?:why does|why do|why is|why are)\b.*\b(?:believe|think|argue|claim|hold|maintain)\b', 400),
    ]

    for pattern, min_words in complex_question_indicators:
        if re.search(pattern, question_lower):
            override_length = max(override_length, min_words)
            override_quotes = max(override_quotes, 3)
            break

    override_length = min(override_length, 10000)
    override_quotes = min(override_quotes, 100)

    if override_length != default_length or override_quotes != default_quotes:
        print(f"⚡ EXPLICIT REQUIREMENTS DETECTED: Overriding settings")
        print(f"   Length: {default_length} → {override_length} words")
        print(f"   Quotes: {default_quotes} → {override_quotes}")

    return override_length, override_quotes

app = Flask(__name__)
app.secret_key = os.environ.get('SESSION_SECRET', os.urandom(24))

# Google OAuth login (see google_auth.py / Replit flask_google_oauth integration)
from google_auth import google_auth, ensure_users_table, is_configured as google_login_configured
app.register_blueprint(google_auth)
ensure_users_table()

print("Configuring semantic search systems (lazy-loaded on first request)...")

databases = {
    'kuczynski': SemanticSearch(
        'data/KUCZYNSKI_PHILOSOPHICAL_DATABASE_v42_WITH_BATCH11.json', 
        'data/kuczynski_embeddings.npy'
    ),
    'freud': SemanticSearch(
        'data/FREUD_DATABASE_UNIFIED.json', 
        'data/freud_unified_embeddings.pkl'
    ),
    'jung': SemanticSearch(
        'data/JUNG_DATABASE.json', 
        'data/jung_embeddings.pkl'
    ),
    'hume': SemanticSearch(
        'data/HUME_DATABASE.json', 
        'data/hume_embeddings.pkl'
    ),
    'nietzsche': SemanticSearch(
        'data/NIETZSCHE_DATABASE.json', 
        'data/nietzsche_embeddings.pkl'
    ),
    'bergler': SemanticSearch(
        'data/BERGLER_DATABASE.json', 
        'data/bergler_embeddings.pkl'
    )
}

print(f"Available databases: {', '.join(databases.keys())} (will load on first use)")

freud_engine = get_freud_engine()
jung_engine = get_jung_engine()
kuczynski_engine = get_kuczynski_engine()
hume_engine = get_hume_engine()
nietzsche_engine = get_nietzsche_engine()
bergler_engine = get_bergler_engine()
print("Inference engines configured (lazy-loaded on first use)")

DATABASE_URL = os.environ.get('DATABASE_URL')

anthropic_client = None
openai_client = None
deepseek_client = None
perplexity_client = None
grok_client = None
venice_client = None

try:
    # Prefer the user's own ANTHROPIC_API_KEY when provided.
    # Falls back to Replit AI Integrations (no user API key required) if no own key is set.
    AI_INTEGRATIONS_ANTHROPIC_API_KEY = os.environ.get("AI_INTEGRATIONS_ANTHROPIC_API_KEY")
    AI_INTEGRATIONS_ANTHROPIC_BASE_URL = os.environ.get("AI_INTEGRATIONS_ANTHROPIC_BASE_URL")
    ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
    if ANTHROPIC_API_KEY and Anthropic:
        anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)
        print("✓ Anthropic client initialized (user API key)")
    elif AI_INTEGRATIONS_ANTHROPIC_API_KEY and AI_INTEGRATIONS_ANTHROPIC_BASE_URL and Anthropic:
        anthropic_client = Anthropic(
            api_key=AI_INTEGRATIONS_ANTHROPIC_API_KEY,
            base_url=AI_INTEGRATIONS_ANTHROPIC_BASE_URL,
        )
        print("✓ Anthropic client initialized (Replit AI Integrations)")
except Exception as e:
    print(f"✗ Could not initialize Anthropic: {e}")

try:
    OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
    if OPENAI_API_KEY and OpenAI:
        openai_client = OpenAI(api_key=OPENAI_API_KEY)
        print("✓ OpenAI client initialized")
except Exception as e:
    print(f"✗ Could not initialize OpenAI: {e}")

try:
    DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
    if DEEPSEEK_API_KEY and OpenAI:
        deepseek_client = OpenAI(
            api_key=DEEPSEEK_API_KEY,
            base_url="https://api.deepseek.com"
        )
        print("✓ DeepSeek client initialized")
except Exception as e:
    print(f"✗ Could not initialize DeepSeek: {e}")

try:
    VENICE_API_KEY = os.environ.get("VENICE_API_KEY")
    if VENICE_API_KEY and OpenAI:
        venice_client = OpenAI(
            api_key=VENICE_API_KEY,
            base_url="https://api.venice.ai/api/v1"
        )
        print("✓ Venice client initialized")
except Exception as e:
    print(f"✗ Could not initialize Venice: {e}")

try:
    PERPLEXITY_API_KEY = os.environ.get("PERPLEXITY_API_KEY")
    if PERPLEXITY_API_KEY and OpenAI:
        perplexity_client = OpenAI(
            api_key=PERPLEXITY_API_KEY,
            base_url="https://api.perplexity.ai"
        )
        print("✓ Perplexity client initialized")
except Exception as e:
    print(f"✗ Could not initialize Perplexity: {e}")

grok_client = None
try:
    XAI_API_KEY = os.environ.get("XAI_API_KEY")
    if XAI_API_KEY and OpenAI:
        grok_client = OpenAI(
            api_key=XAI_API_KEY,
            base_url="https://api.x.ai/v1"
        )
        print("✓ Grok (xAI) client initialized")
except Exception as e:
    print(f"✗ Could not initialize Grok: {e}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/download-embeddings')
def download_embeddings():
    """Temporary endpoint to download embeddings archive"""
    from flask import send_file
    embeddings_path = 'embeddings.tar.gz'
    if os.path.exists(embeddings_path):
        return send_file(embeddings_path, as_attachment=True, download_name='embeddings.tar.gz')
    return "File not found", 404

def get_db_position_counts():
    """Get position counts from PostgreSQL database"""
    try:
        import psycopg2
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT thinker, COUNT(*) FROM positions GROUP BY thinker")
        counts = {row[0]: row[1] for row in cur.fetchall()}
        cur.close()
        conn.close()
        return counts
    except Exception as e:
        print(f"Error getting DB counts: {e}")
        return {
            'kuczynski': 7434,
            'freud': 9618,
            'jung': 3244,
            'hume': 1273,
            'nietzsche': 1261,
            'bergler': 107
        }

def search_positions_postgres(question, thinker, top_k=15):
    """Search positions from PostgreSQL using full-text search"""
    try:
        import psycopg2
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()

        words = re.findall(r'\b\w{3,}\b', question.lower())
        stop_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her', 'was', 'one', 'our', 'out', 'his', 'has', 'have', 'with', 'they', 'this', 'from', 'that', 'what', 'how', 'why', 'when', 'where', 'which', 'who', 'does', 'your'}
        keywords = [w for w in words if w not in stop_words][:10]

        if not keywords:
            cur.close()
            conn.close()
            return []

        tsquery = ' | '.join(keywords)

        query = """
            SELECT id, thinker, position, topic,
                   ts_rank(to_tsvector('english', position || ' ' || COALESCE(topic, '')), to_tsquery('english', %s)) as rank
            FROM positions
            WHERE thinker = %s
              AND to_tsvector('english', position || ' ' || COALESCE(topic, '')) @@ to_tsquery('english', %s)
            ORDER BY rank DESC
            LIMIT %s
        """

        cur.execute(query, (tsquery, thinker, tsquery, top_k))
        rows = cur.fetchall()

        positions = []
        for row in rows:
            positions.append({
                'position_id': f"PG-{row[0]}",
                'text': row[2],
                'title': row[3] or '',
                'domain': row[3] or 'General',
                'similarity': float(row[4]) if row[4] else 0.5,
                'source': ['PostgreSQL']
            })

        cur.close()
        conn.close()

        print(f"📊 PostgreSQL: Found {len(positions)} positions for '{thinker}' (keywords: {keywords[:5]})")
        return positions

    except Exception as e:
        print(f"PostgreSQL search error: {e}")
        return []

def search_quotes_postgres(question, thinker, top_k=5):
    """Search quotes from PostgreSQL using full-text search"""
    try:
        import psycopg2
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()

        words = re.findall(r'\b\w{3,}\b', question.lower())
        stop_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her', 'was', 'one', 'our', 'out', 'his', 'has', 'have', 'with', 'they', 'this', 'from', 'that', 'what', 'how', 'why', 'when', 'where', 'which', 'who', 'does', 'your'}
        keywords = [w for w in words if w not in stop_words][:10]

        if not keywords:
            cur.close()
            conn.close()
            return []

        tsquery = ' | '.join(keywords)

        query = """
            SELECT id, thinker, quote, topic,
                   ts_rank(to_tsvector('english', quote || ' ' || COALESCE(topic, '')), to_tsquery('english', %s)) as rank
            FROM quotes
            WHERE thinker = %s
              AND to_tsvector('english', quote || ' ' || COALESCE(topic, '')) @@ to_tsquery('english', %s)
            ORDER BY rank DESC
            LIMIT %s
        """

        cur.execute(query, (tsquery, thinker, tsquery, top_k))
        rows = cur.fetchall()

        quotes = []
        for row in rows:
            quotes.append({
                'position_id': f"QUOTE-{row[0]}",
                'text': row[2],
                'title': row[3] or '',
                'domain': 'Quote',
                'similarity': float(row[4]) if row[4] else 0.5,
                'source': ['PostgreSQL Quotes']
            })

        cur.close()
        conn.close()

        print(f"💬 PostgreSQL Quotes: Found {len(quotes)} quotes for '{thinker}'")
        return quotes

    except Exception as e:
        print(f"PostgreSQL quotes search error: {e}")
        return []

def merge_position_results(tier1, tier2, max_results=15):
    """Merge and deduplicate position results from two sources"""
    from difflib import SequenceMatcher

    combined = []
    seen_texts = []

    def is_duplicate(text, threshold=0.85):
        text_lower = text.lower()[:200]
        for seen in seen_texts:
            if SequenceMatcher(None, text_lower, seen).ratio() > threshold:
                return True
        return False

    for pos in tier1:
        text = pos.get('text', '')
        if text and not is_duplicate(text):
            combined.append(pos)
            seen_texts.append(text.lower()[:200])

    for pos in tier2:
        text = pos.get('text', '')
        if text and not is_duplicate(text):
            combined.append(pos)
            seen_texts.append(text.lower()[:200])

    combined.sort(key=lambda x: x.get('similarity', 0), reverse=True)

    return combined[:max_results]

def search_text_chunks_rag(question, thinker, top_k=5):
    """RAG: Search text_chunks table for relevant source text chunks using PostgreSQL full-text search"""
    try:
        import psycopg2
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()

        words = re.findall(r'\b\w{3,}\b', question.lower())
        stop_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'her', 'was', 'one', 'our', 'out', 'his', 'has', 'have', 'with', 'they', 'this', 'from', 'that', 'what', 'how', 'why', 'when', 'where', 'which', 'who', 'does', 'your'}
        keywords = [w for w in words if w not in stop_words][:10]

        if not keywords:
            cur.close()
            conn.close()
            return []

        tsquery = ' | '.join(keywords)

        query = """
            SELECT source_file, chunk_text, chunk_index,
                   ts_rank(to_tsvector('english', chunk_text), to_tsquery('english', %s)) as rank
            FROM texts
            WHERE thinker = %s
              AND to_tsvector('english', chunk_text) @@ to_tsquery('english', %s)
            ORDER BY rank DESC
            LIMIT %s
        """

        cur.execute(query, (tsquery, thinker, tsquery, top_k))
        rows = cur.fetchall()

        chunks = []
        for row in rows:
            chunks.append({
                'source_file': row[0],
                'chunk_text': row[1],
                'chunk_index': row[2],
                'relevance': float(row[3])
            })

        cur.close()
        conn.close()

        print(f"📚 RAG: Found {len(chunks)} relevant text chunks for '{thinker}' (keywords: {keywords[:5]})")
        return chunks

    except Exception as e:
        print(f"RAG search error: {e}")
        return []

@app.route('/api/databases', methods=['GET'])
def get_databases():
    """Return available databases (Jung hidden from UI but logic preserved)"""
    display_names = {
        'kuczynski': 'ZHI',
        'freud': 'Freud',
        'jung': 'Jung',
        'hume': 'Hume',
        'nietzsche': 'Nietzsche',
        'bergler': 'Bergler'
    }
    position_counts = get_db_position_counts()
    hidden_databases = set()
    db_list = []
    for db_id, db_obj in databases.items():
        if db_id in hidden_databases:
            continue
        db_list.append({
            'id': db_id,
            'name': display_names.get(db_id, db_id.capitalize()),
            'count': position_counts.get(db_id, 0)
        })
    return jsonify({'databases': db_list})

@app.route('/api/positions/search', methods=['GET'])
def search_positions_db():
    """Search positions directly from PostgreSQL database"""
    try:
        import psycopg2
        thinker = request.args.get('thinker', '')
        search_term = request.args.get('q', '')
        limit = min(int(request.args.get('limit', 50)), 200)
        offset = int(request.args.get('offset', 0))

        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()

        query = "SELECT id, thinker, position, topic, created_at FROM positions WHERE 1=1"
        params = []

        if thinker:
            query += " AND thinker = %s"
            params.append(thinker)

        if search_term:
            query += " AND (position ILIKE %s OR topic ILIKE %s)"
            params.extend([f'%{search_term}%', f'%{search_term}%'])

        query += " ORDER BY id LIMIT %s OFFSET %s"
        params.extend([limit, offset])

        cur.execute(query, params)
        rows = cur.fetchall()

        positions = []
        for row in rows:
            positions.append({
                'id': row[0],
                'thinker': row[1],
                'position': row[2],
                'topic': row[3],
                'created_at': str(row[4]) if row[4] else None
            })

        cur.close()
        conn.close()

        return jsonify({'positions': positions, 'count': len(positions)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/providers', methods=['GET'])
def get_providers():
    """Return available AI providers in order ZHI 1-5, with ZHI 4 (DeepSeek) as default"""
    providers = []
    # ZHI 1 - Anthropic
    if anthropic_client:
        providers.append({'id': 'anthropic', 'name': 'ZHI 1', 'models': ['claude-sonnet-4-5-20250929', 'claude-opus-4-5']})
    # ZHI 2 - OpenAI
    if openai_client:
        providers.append({'id': 'openai', 'name': 'ZHI 2', 'models': ['gpt-4o', 'gpt-4o-mini', 'o1', 'o1-mini']})
    # ZHI 3 - Grok
    if grok_client:
        providers.append({'id': 'grok', 'name': 'ZHI 3', 'models': ['grok-4', 'grok-3-beta', 'grok-3-mini-beta', 'grok-code-fast-1']})
    # ZHI 4 - DeepSeek (DEFAULT)
    if deepseek_client:
        providers.append({'id': 'deepseek', 'name': 'ZHI 4', 'models': ['deepseek-chat', 'deepseek-reasoner'], 'default': True})
    # ZHI 5 - Perplexity
    if perplexity_client:
        providers.append({'id': 'perplexity', 'name': 'ZHI 5', 'models': ['sonar-pro', 'sonar', 'sonar-reasoning']})
    # ZHI 6 - Venice
    if venice_client:
        providers.append({'id': 'venice', 'name': 'ZHI 6', 'models': ['llama-3.3-70b', 'venice-uncensored', 'qwen3-235b', 'llama-3.1-405b', 'dolphin-2.9.2-qwen2-72b']})
    return jsonify({'providers': providers})

def get_fallback_providers(primary):
    """Get ordered list of fallback providers after primary fails"""
    all_providers = []
    if anthropic_client:
        all_providers.append('anthropic')
    if openai_client:
        all_providers.append('openai')
    if grok_client:
        all_providers.append('grok')
    if deepseek_client:
        all_providers.append('deepseek')
    if venice_client:
        all_providers.append('venice')
    if perplexity_client:
        all_providers.append('perplexity')
    return [p for p in all_providers if p != primary]

@app.route('/api/ask', methods=['POST'])
def ask():
    """Handle user question with streaming response"""
    try:
        data = request.json
        question = data.get('question', '')
        provider = data.get('provider', 'deepseek')
        model = data.get('model', '')
        database = data.get('database', 'freud')
        enhanced_mode = data.get('enhanced_mode', False)
        try:
            creativity_level = min(max(int(data.get('creativity_level', 10)), 1), 20)
        except (TypeError, ValueError):
            creativity_level = 10
        gen_temperature = creativity_temperature(creativity_level)
        answer_length = min(max(data.get('answer_length', 250), 100), 8000)
        quote_count = min(max(data.get('quote_count', 5), 1), 50)
        response_mode = data.get('response_mode', 'standard')
        data_source = data.get('data_source', 'combined')
        elevenlabs_mode = bool(data.get('elevenlabs_mode', False))

        if response_mode not in ('dialogue', 'conversation'):
            answer_length, quote_count = detect_explicit_requirements(question, answer_length, quote_count)

        if answer_length >= 2000:
            base_top_k = 30
            rag_top_k = 20
        elif answer_length >= 1500:
            base_top_k = 25
            rag_top_k = 15
        elif answer_length >= 1000:
            base_top_k = 20
            rag_top_k = 10
        elif answer_length >= 500:
            base_top_k = 15
            rag_top_k = 8
        else:
            base_top_k = 10
            rag_top_k = 5

        print(f"Received question: {question}")
        print(f"Provider: {provider}, Model: {model}, Database: {database}, Enhanced Mode: {enhanced_mode}")
        print(f"Answer Length: {answer_length} words, Quote Count: {quote_count}, Response Mode: {response_mode}, Data Source: {data_source}, ElevenLabs: {elevenlabs_mode}")
        print(f"Retrieval scaling: base_top_k={base_top_k}, rag_top_k={rag_top_k}")

        if not question:
            return jsonify({'error': 'No question provided'}), 400

        if database not in databases:
            return jsonify({'error': f'Database "{database}" not available'}), 400

        searcher = databases[database]

        # Default to Freud if database parameter not provided
        if not database or database not in databases:
            database = 'freud' if 'freud' in databases else list(databases.keys())[0]
            searcher = databases[database]

        print(f"Searching {database} database for relevant positions...")
        try:
            # CANONICAL QUERY MAPPING: Force-include key positions for well-defined questions
            canonical_position_ids = []
            question_lower = question.lower()

            if database == 'kuczynski':
                # Proposition composition questions
                if any(term in question_lower for term in ['composition of', 'decompos', 'consist of', 'constituents of']) and 'proposition' in question_lower:
                    canonical_position_ids = ['LMCC-323', 'LMCC-324', 'ANALPHIL-062', 'ANALPHIL-142', 'LSPM-095', 'CONCAUS-084', 'CONCAUS-088']
                    print(f"⚡ CANONICAL QUERY DETECTED: proposition composition → forcing {len(canonical_position_ids)} key positions")
                # What are propositions questions  
                elif 'what are proposition' in question_lower or 'what is a proposition' in question_lower:
                    canonical_position_ids = ['LMCC-323', 'ANALPHIL-062', 'ANALPHIL-142', 'MMSE-014', 'MMSE-047']
                    print(f"⚡ CANONICAL QUERY DETECTED: what are propositions → forcing {len(canonical_position_ids)} key positions")
                # Truth and instantiation questions
                elif 'truth' in question_lower and ('instantiat' in question_lower or 'property' in question_lower):
                    canonical_position_ids = ['LMCC-323', 'ANALPHIL-062', 'ANALPHIL-142']
                    print(f"⚡ CANONICAL QUERY DETECTED: truth/instantiation → forcing {len(canonical_position_ids)} key positions")

            # Get search results based on data_source setting - FIX: use scaled top_k values
            use_rag = False
            retrieval_metadata = None
            if data_source == 'classic':
                # Classic: Use only existing JSON/embeddings semantic search with full metadata
                search_result = searcher.search(question, top_k=base_top_k, return_metadata=True)
                relevant_positions = search_result['results']
                retrieval_metadata = search_result['metadata']
                print(f"📚 Using CLASSIC search (semantic embeddings) - top_k={base_top_k}")
                print(f"   📊 SCANNED {retrieval_metadata['total_positions_scanned']} positions, {retrieval_metadata['positions_above_threshold']} above threshold")
                print(f"   📈 Max similarity: {retrieval_metadata['max_similarity']:.3f}, Mean: {retrieval_metadata['mean_similarity']:.3f}")
            elif data_source == 'newdb':
                # New DB: Use only PostgreSQL tables
                relevant_positions = search_positions_postgres(question, database, top_k=base_top_k)
                quote_results = search_quotes_postgres(question, database, top_k=rag_top_k)
                relevant_positions = merge_position_results(relevant_positions, quote_results, max_results=base_top_k)
                use_rag = True
                # Generate metadata for PostgreSQL search (note: PG full-text doesn't compute cosine similarity)
                max_rank = max([p.get('rank', p.get('similarity', 0)) for p in relevant_positions]) if relevant_positions else 0
                mean_rank = sum([p.get('rank', p.get('similarity', 0)) for p in relevant_positions]) / len(relevant_positions) if relevant_positions else 0
                retrieval_metadata = {
                    'total_positions_scanned': 'N/A (full-text)',
                    'positions_above_threshold': len(relevant_positions),
                    'threshold': 'N/A',
                    'max_similarity': max_rank,
                    'mean_similarity': mean_rank,
                    'top_k_used': base_top_k,
                    'domains_in_results': list(set(p.get('domain', 'Unknown') for p in relevant_positions)),
                    'search_type': 'postgresql_fulltext',
                    'note': 'PostgreSQL full-text search uses ts_rank, not cosine similarity'
                }
                print(f"📊 Using NEW DB search (PostgreSQL full-text) - top_k={base_top_k}")
            elif data_source == 'combined':
                # Combined: Use both systems and merge results with full metadata
                search_result = searcher.search(question, top_k=base_top_k, return_metadata=True)
                tier1_results = search_result['results']
                retrieval_metadata = search_result['metadata']
                retrieval_metadata['search_type'] = 'combined_semantic_postgresql'
                tier2_results = search_positions_postgres(question, database, top_k=base_top_k)
                quote_results = search_quotes_postgres(question, database, top_k=rag_top_k)
                tier2_combined = merge_position_results(tier2_results, quote_results, max_results=base_top_k)
                relevant_positions = merge_position_results(tier1_results, tier2_combined, max_results=base_top_k + 5)
                use_rag = True
                print(f"🔀 Using COMBINED search (embeddings + PostgreSQL) - top_k={base_top_k}")
                print(f"   📊 SCANNED {retrieval_metadata['total_positions_scanned']} positions, {retrieval_metadata['positions_above_threshold']} above threshold")
                print(f"   📈 Max similarity: {retrieval_metadata['max_similarity']:.3f}, Mean: {retrieval_metadata['mean_similarity']:.3f}")
            else:
                # Default to combined (not classic!) with full metadata
                search_result = searcher.search(question, top_k=base_top_k, return_metadata=True)
                tier1_results = search_result['results']
                retrieval_metadata = search_result['metadata']
                retrieval_metadata['search_type'] = 'combined_semantic_postgresql'
                tier2_results = search_positions_postgres(question, database, top_k=base_top_k)
                relevant_positions = merge_position_results(tier1_results, tier2_results, max_results=base_top_k)
                use_rag = True
                print(f"   📊 SCANNED {retrieval_metadata['total_positions_scanned']} positions, {retrieval_metadata['positions_above_threshold']} above threshold")

            # Force-include canonical positions at the top if they exist
            if canonical_position_ids:
                # Get the canonical positions from database
                canonical_positions = []
                for pos in searcher.positions:
                    if pos.get('position_id') in canonical_position_ids:
                        # Add with high similarity to ensure they appear first
                        pos_copy = pos.copy()
                        pos_copy['similarity'] = 0.99  # Force to top
                        canonical_positions.append(pos_copy)

                # Remove duplicates from search results
                search_ids = {p.get('position_id') for p in relevant_positions}
                canonical_positions = [p for p in canonical_positions if p.get('position_id') not in search_ids or p.get('position_id') in canonical_position_ids]

                # Prepend canonical positions
                relevant_positions = canonical_positions + [p for p in relevant_positions if p.get('position_id') not in canonical_position_ids]
                print(f"   Injected {len(canonical_positions)} canonical positions at top")

            # Expand results with argument context (related/adjacent positions)
            context_expanded_count = 0
            if data_source in ['classic', 'combined', None] and hasattr(searcher, 'expand_with_context'):
                original_count = len(relevant_positions)
                relevant_positions = searcher.expand_with_context(relevant_positions, max_context=2)
                context_expanded_count = len(relevant_positions) - original_count
                if context_expanded_count > 0:
                    print(f"   📖 Expanded {original_count} → {len(relevant_positions)} with argument context")
                    # Update retrieval metadata to reflect expanded positions
                    if retrieval_metadata:
                        retrieval_metadata['context_positions_added'] = context_expanded_count
                        retrieval_metadata['total_positions_used'] = len(relevant_positions)

            print(f"Found {len(relevant_positions)} relevant positions")

            # Detect low-relevance searches (external knowledge mode)
            max_similarity = max([p.get('similarity', 0) for p in relevant_positions]) if relevant_positions else 0
            low_relevance = max_similarity < 0.40

            if low_relevance:
                print(f"⚠️  LOW RELEVANCE DETECTED (max similarity: {max_similarity:.3f})")
                print("   Activating External Knowledge Assimilation mode...")
            else:
                print(f"✓ Good relevance (max similarity: {max_similarity:.3f})")

        except Exception as e:
            print(f"ERROR in search: {str(e)} - CONTINUING WITH LLM ANYWAY")
            import traceback
            traceback.print_exc()
            # DON'T CRASH - continue with empty positions, LLM will still answer
            relevant_positions = []
            rag_chunks = []
            low_relevance = True  # Trigger external knowledge mode

        # INFERENCE ENGINE: Deduce rules based on thinker
        deduced_rules = []
        if freud_engine and database.startswith('freud'):
            try:
                print("🧠 Activating Freud inference engine...")
                deduced_rules = freud_engine.deduce(question, max_rules=15)
                print(f"✓ Fired {len(deduced_rules)} inference rules")
                for i, rule in enumerate(deduced_rules[:5], 1):
                    print(f"   {i}. [{rule.get('year', '?')}] {rule.get('id', '?')}: {rule.get('conclusion', '')[:80]}...")
            except Exception as e:
                print(f"⚠️  Freud inference engine error: {e}")
        elif kuczynski_engine and database == 'kuczynski':
            try:
                print("🧠 Activating Kuczynski inference engine...")
                deduced_rules = kuczynski_engine.deduce(question, max_rules=15)
                print(f"✓ Fired {len(deduced_rules)} inference rules")
                for i, rule in enumerate(deduced_rules[:5], 1):
                    print(f"   {i}. [{rule.get('year', '?')}] {rule.get('id', '?')}: {rule.get('conclusion', '')[:80]}...")
            except Exception as e:
                print(f"⚠️  Kuczynski inference engine error: {e}")
        elif jung_engine and database == 'jung':
            try:
                print("🧠 Activating Jung inference engine...")
                deduced_rules = jung_engine.deduce(question, max_rules=15)
                print(f"✓ Fired {len(deduced_rules)} inference rules")
                for i, rule in enumerate(deduced_rules[:5], 1):
                    print(f"   {i}. [{rule.get('year', '?')}] {rule.get('id', '?')}: {rule.get('conclusion', '')[:80]}...")
            except Exception as e:
                print(f"⚠️  Jung inference engine error: {e}")
        elif hume_engine and database == 'hume':
            try:
                print("🧠 Activating Hume inference engine...")
                deduced_rules = hume_engine.deduce(question, max_rules=15)
                print(f"✓ Fired {len(deduced_rules)} inference rules")
                for i, rule in enumerate(deduced_rules[:5], 1):
                    print(f"   {i}. [{rule.get('year', '?')}] {rule.get('id', '?')}: {rule.get('conclusion', '')[:80]}...")
            except Exception as e:
                print(f"⚠️  Hume inference engine error: {e}")
        elif nietzsche_engine and database == 'nietzsche':
            try:
                print("🧠 Activating Nietzsche inference engine...")
                deduced_rules = nietzsche_engine.deduce(question, max_rules=15)
                print(f"✓ Fired {len(deduced_rules)} inference rules")
                for i, rule in enumerate(deduced_rules[:5], 1):
                    print(f"   {i}. [{rule.get('year', '?')}] {rule.get('id', '?')}: {rule.get('conclusion', '')[:80]}...")
            except Exception as e:
                print(f"⚠️  Nietzsche inference engine error: {e}")
        elif bergler_engine and database == 'bergler':
            try:
                print("🧠 Activating Bergler inference engine...")
                deduced_rules = bergler_engine.deduce(question, max_rules=15)
                print(f"✓ Fired {len(deduced_rules)} inference rules")
                for i, rule in enumerate(deduced_rules[:5], 1):
                    print(f"   {i}. [{rule.get('year', '?')}] {rule.get('id', '?')}: {rule.get('conclusion', '')[:80]}...")
            except Exception as e:
                print(f"⚠️  Bergler inference engine error: {e}")

        # Store deduced rules in session for debug endpoint
        session['last_deduced_rules'] = deduced_rules

        # FIX: ALWAYS search text chunks for richer context (essential for long responses)
        rag_chunks = search_text_chunks_rag(question, database, top_k=rag_top_k)
        if rag_chunks:
            print(f"📖 RAG: Loaded {len(rag_chunks)} source text chunks for grounding")

        # Get or create conversation ID
        if 'conversation_id' not in session:
            session['conversation_id'] = conversation_manager.get_conversation_id()
        conversation_id = session['conversation_id']

        # Get conversation history (filtered to current database only to prevent cross-contamination)
        conversation_history = conversation_manager.format_history_for_prompt(conversation_id, max_recent=5, current_database=database)

        # Track the full answer for storage after streaming
        full_answer = []

        # Create question hash for resumable response tracking
        import hashlib
        question_hash = hashlib.md5(f"{question}:{database}:{answer_length}".encode()).hexdigest()[:16]
        last_saved_word_count = [0]  # Track when we last saved to DB
        SAVE_INTERVAL = 100  # Save every 100 words

        def get_db_connection():
            """Get PostgreSQL connection for response progress tracking."""
            import psycopg2
            return psycopg2.connect(os.environ.get('DATABASE_URL'))

        def check_incomplete_response():
            """Check if there's an incomplete response we can resume from."""
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    SELECT accumulated_response, word_count, target_word_count
                    FROM response_progress
                    WHERE conversation_id = %s AND question_hash = %s AND is_complete = FALSE
                """, (conversation_id, question_hash))
                row = cur.fetchone()
                cur.close()
                conn.close()
                if row and row[1] > 50:  # Only resume if we have at least 50 words
                    return {'text': row[0], 'word_count': row[1], 'target': row[2]}
                return None
            except Exception as e:
                print(f"Error checking incomplete response: {e}")
                return None

        def save_response_progress(text, word_count, target_word_count, is_complete=False):
            """Save response progress to database."""
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    INSERT INTO response_progress (conversation_id, question_hash, accumulated_response, word_count, target_word_count, is_complete, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, NOW())
                    ON CONFLICT (conversation_id, question_hash)
                    DO UPDATE SET accumulated_response = %s, word_count = %s, is_complete = %s, updated_at = NOW()
                """, (conversation_id, question_hash, text, word_count, target_word_count, is_complete, text, word_count, is_complete))
                conn.commit()
                cur.close()
                conn.close()
                print(f"💾 Saved progress: {word_count} words (complete={is_complete})")
            except Exception as e:
                print(f"Error saving response progress: {e}")

        def delete_completed_response():
            """Delete completed response from progress table."""
            try:
                conn = get_db_connection()
                cur = conn.cursor()
                cur.execute("""
                    DELETE FROM response_progress
                    WHERE conversation_id = %s AND question_hash = %s
                """, (conversation_id, question_hash))
                conn.commit()
                cur.close()
                conn.close()
                print(f"🗑️ Cleaned up completed response progress")
            except Exception as e:
                print(f"Error deleting response progress: {e}")

        def maybe_save_progress():
            """Save progress every SAVE_INTERVAL words."""
            current_word_count = len(accumulated_text[0].split())
            if current_word_count - last_saved_word_count[0] >= SAVE_INTERVAL:
                last_saved_word_count[0] = current_word_count
                save_response_progress(accumulated_text[0], current_word_count, answer_length, is_complete=False)

        # Word counter for rate limit pauses - accumulate text and count complete words
        accumulated_text = ['']
        last_pause_word_count = [0]
        last_heartbeat_time = [time.time()]
        PAUSE_THRESHOLD = 150  # Pause every 150 words
        PAUSE_DURATION = 5   # Pause for 5 seconds
        HEARTBEAT_INTERVAL = 15  # Send heartbeat every 15 seconds to keep connection alive

        def get_heartbeat_if_needed():
            """Check if a heartbeat is needed (every 15 seconds) to keep connection alive on Render/proxies."""
            current_time = time.time()
            if current_time - last_heartbeat_time[0] >= HEARTBEAT_INTERVAL:
                last_heartbeat_time[0] = current_time
                return f": keepalive {int(current_time)}\n\n"
            return None

        def count_words_and_pause(text):
            """Accumulate streaming text and check if pause needed. GENERATOR that yields heartbeats during pause."""
            accumulated_text[0] += text
            last_heartbeat_time[0] = time.time()  # Reset heartbeat timer on actual content

            # Count complete words (split by whitespace)
            current_word_count = len(accumulated_text[0].split())
            words_since_last_pause = current_word_count - last_pause_word_count[0]

            # Check if we've crossed another threshold
            if words_since_last_pause >= PAUSE_THRESHOLD:
                last_pause_word_count[0] = current_word_count
                print(f"⏸️ Rate limit pause at {current_word_count} words (sleeping {PAUSE_DURATION}s with heartbeats)...")
                # Sleep in small intervals using gevent.sleep (cooperative), YIELDING heartbeats IMMEDIATELY
                for i in range(PAUSE_DURATION * 2):  # Every 0.5 seconds
                    gevent.sleep(0.5)  # Cooperative sleep - allows heartbeat greenlet to run
                    # SSE comment that keeps connection alive - YIELD immediately, don't collect
                    yield f": heartbeat {current_word_count}\n\n"
                print(f"▶️ Resuming after pause...")
                last_heartbeat_time[0] = time.time()  # Reset after pause

        def try_provider(prov, prompt_text, sys_prompt):
            """
            Try streaming from a provider. Returns generator that yields SSE events.
            Raises exception on failure.
            """
            if prov == 'anthropic':
                if not anthropic_client:
                    raise Exception("Anthropic not configured")
                model_name = model if model and prov == provider else "claude-sonnet-4-5-20250929"
                if model_name not in ("claude-sonnet-4-5-20250929", "claude-opus-4-5"):
                    model_name = "claude-sonnet-4-5-20250929"
                print(f"🔄 Trying Anthropic ({model_name})...")
                with anthropic_client.messages.stream(
                    model=model_name,
                    max_tokens=16000,
                    temperature=gen_temperature,
                    system=sys_prompt,
                    messages=[{"role": "user", "content": prompt_text}]
                ) as stream:
                    for text in stream.text_stream:
                        # Check if we need a heartbeat before sending content
                        hb = get_heartbeat_if_needed()
                        if hb:
                            yield hb
                        full_answer.append(text)
                        yield f"data: {json.dumps({'type': 'token', 'data': text})}\n\n"
                        for hb in count_words_and_pause(text):
                            yield hb
                print(f"✓ Anthropic completed: {len(full_answer)} tokens, {len(accumulated_text[0].split())} words")

            elif prov == 'openai':
                if not openai_client:
                    raise Exception("OpenAI not configured")
                model_name = model if model and prov == provider else "gpt-4o"
                print(f"🔄 Trying OpenAI ({model_name})...")
                stream = openai_client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": sys_prompt},
                        {"role": "user", "content": prompt_text}
                    ],
                    stream=True,
                    max_tokens=16000,
                    temperature=gen_temperature,
                    timeout=180.0
                )
                for chunk in stream:
                    # Check if we need a heartbeat
                    hb = get_heartbeat_if_needed()
                    if hb:
                        yield hb
                    if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_answer.append(content)
                        yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                        for hb in count_words_and_pause(content):
                            yield hb
                    if chunk.choices and chunk.choices[0].finish_reason:
                        break
                print(f"✓ OpenAI completed: {len(full_answer)} tokens, {len(accumulated_text[0].split())} words")

            elif prov == 'grok':
                if not grok_client:
                    raise Exception("Grok not configured")
                model_name = model if model and prov == provider else "grok-4"
                print(f"🔄 Trying Grok ({model_name})...")
                stream = grok_client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": sys_prompt},
                        {"role": "user", "content": prompt_text}
                    ],
                    stream=True,
                    max_tokens=16000,
                    temperature=gen_temperature,
                    timeout=180.0
                )
                for chunk in stream:
                    # Check if we need a heartbeat
                    hb = get_heartbeat_if_needed()
                    if hb:
                        yield hb
                    if chunk.choices and len(chunk.choices) > 0:
                        if chunk.choices[0].delta and chunk.choices[0].delta.content:
                            content = chunk.choices[0].delta.content
                            full_answer.append(content)
                            yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                            for hb in count_words_and_pause(content):
                                yield hb
                        if chunk.choices[0].finish_reason:
                            break
                print(f"✓ Grok completed: {len(full_answer)} tokens, {len(accumulated_text[0].split())} words")

            elif prov == 'deepseek':
                if not deepseek_client:
                    raise Exception("DeepSeek not configured")
                model_name = model if model and prov == provider else "deepseek-chat"
                print(f"🔄 Trying DeepSeek ({model_name})...")
                stream = deepseek_client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": sys_prompt},
                        {"role": "user", "content": prompt_text}
                    ],
                    stream=True,
                    max_tokens=8192,
                    temperature=gen_temperature,
                    timeout=300.0
                )
                for chunk in stream:
                    # Check if we need a heartbeat
                    hb = get_heartbeat_if_needed()
                    if hb:
                        yield hb
                    if chunk.choices and len(chunk.choices) > 0:
                        if chunk.choices[0].delta and chunk.choices[0].delta.content:
                            content = chunk.choices[0].delta.content
                            full_answer.append(content)
                            yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                            for hb in count_words_and_pause(content):
                                yield hb
                        if chunk.choices[0].finish_reason:
                            break
                print(f"✓ DeepSeek completed: {len(full_answer)} tokens, {len(accumulated_text[0].split())} words")

            elif prov == 'venice':
                if not venice_client:
                    raise Exception("Venice not configured")
                model_name = model if model and prov == provider else "llama-3.3-70b"
                print(f"🔄 Trying Venice ({model_name})...")
                stream = venice_client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": sys_prompt},
                        {"role": "user", "content": prompt_text}
                    ],
                    stream=True,
                    max_tokens=8192,
                    temperature=gen_temperature,
                    timeout=300.0
                )
                for chunk in stream:
                    hb = get_heartbeat_if_needed()
                    if hb:
                        yield hb
                    if chunk.choices and len(chunk.choices) > 0:
                        if chunk.choices[0].delta and chunk.choices[0].delta.content:
                            content = chunk.choices[0].delta.content
                            full_answer.append(content)
                            yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                            for hb in count_words_and_pause(content):
                                yield hb
                        if chunk.choices[0].finish_reason:
                            break
                print(f"✓ Venice completed: {len(full_answer)} tokens, {len(accumulated_text[0].split())} words")

            elif prov == 'perplexity':
                if not perplexity_client:
                    raise Exception("Perplexity not configured")
                model_name = model if model and prov == provider else "sonar-pro"
                print(f"🔄 Trying Perplexity ({model_name})...")
                pplx_thinker = {'freud': 'Sigmund Freud', 'kuczynski': 'J.-M. Kuczynski', 'jung': 'Carl Gustav Jung', 'hume': 'David Hume', 'nietzsche': 'Friedrich Nietzsche'}.get(database, database.capitalize())
                pplx_sys = f"""You are an educational philosophy assistant presenting documented philosophical positions from {pplx_thinker}'s actual writings.
Present the material in first person as is standard in academic philosophy education.
{sys_prompt}
NEVER break character or add meta-commentary about the task."""
                stream = perplexity_client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": pplx_sys},
                        {"role": "user", "content": prompt_text}
                    ],
                    stream=True,
                    max_tokens=16000,
                    temperature=gen_temperature
                )
                for chunk in stream:
                    # Check if we need a heartbeat
                    hb = get_heartbeat_if_needed()
                    if hb:
                        yield hb
                    if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_answer.append(content)
                        yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                        for hb in count_words_and_pause(content):
                            yield hb
                    if chunk.choices and chunk.choices[0].finish_reason:
                        break
                print(f"✓ Perplexity completed: {len(full_answer)} tokens, {len(accumulated_text[0].split())} words")

            else:
                raise Exception(f"Unknown provider: {prov}")

        def generate():
            """Queue-based SSE generator with dedicated heartbeat greenlet for bulletproof connection keepalive."""
            event_queue = gevent_queue.Queue()
            streaming_done = [False]  # Mutable flag for greenlet coordination
            HEARTBEAT_INTERVAL = 3  # Send heartbeat every 3 seconds (very aggressive for Render)

            def heartbeat_worker():
                """Dedicated greenlet that sends heartbeats every 8 seconds until streaming is done."""
                heartbeat_count = 0
                while not streaming_done[0]:
                    gevent.sleep(HEARTBEAT_INTERVAL)
                    if not streaming_done[0]:
                        heartbeat_count += 1
                        event_queue.put(f": keepalive {heartbeat_count}\n\n")
                print(f"Heartbeat worker stopped after {heartbeat_count} heartbeats")

            def content_worker():
                """Greenlet that streams content from AI providers into the queue."""
                try:
                    print("Starting SSE content worker with automatic failover...")

                    # Check for incomplete response to resume
                    incomplete = check_incomplete_response()
                    resumed_text = ""
                    if incomplete:
                        resumed_text = incomplete['text']
                        print(f"🔄 RESUMING from {incomplete['word_count']} words (target: {incomplete['target']})")
                        # Send resumed text as a single chunk first
                        event_queue.put(f"data: {json.dumps({'type': 'resumed', 'data': resumed_text, 'word_count': incomplete['word_count']})}\n\n")
                        # Pre-populate our tracking
                        accumulated_text[0] = resumed_text
                        full_answer.append(resumed_text)
                        last_saved_word_count[0] = incomplete['word_count']

                    sources = [p['position_id'] for p in relevant_positions]
                    positions_data = [{
                        'id': p.get('position_id', ''),
                        'text': (p.get('text') or p.get('thesis', ''))[:300] + ('...' if len(p.get('text') or p.get('thesis', '')) > 300 else ''),
                        'domain': p.get('domain', 'philosophy'),
                        'similarity': p.get('similarity', 0)
                    } for p in relevant_positions]
                    event_queue.put(f"data: {json.dumps({'type': 'sources', 'data': sources, 'positions': positions_data})}\n\n")
                    
                    if retrieval_metadata:
                        event_queue.put(f"data: {json.dumps({'type': 'retrieval_log', 'data': retrieval_metadata})}\n\n")

                    # Build prompt - if resuming, ask to continue from where we left off
                    if resumed_text:
                        continuation_prompt = f"""CONTINUATION REQUEST: The previous response was interrupted at {len(resumed_text.split())} words. Here is what was written so far:

---BEGIN PREVIOUS RESPONSE---
{resumed_text}
---END PREVIOUS RESPONSE---

Please CONTINUE this response from EXACTLY where it left off. Do NOT repeat what was already written. Just continue naturally from the last word. The target is {answer_length} words total.

Original question: {question}"""
                        prompt = continuation_prompt
                    else:
                        prompt = build_prompt(question, relevant_positions, database, conversation_history, enhanced_mode, low_relevance, deduced_rules, answer_length, quote_count, response_mode, rag_chunks, creativity_level)

                    if response_mode == 'dialogue':
                        sys_prompt = f"""You are this philosopher in dialogue. Keep responses BRIEF (2-4 sentences) but USE the retrieved positions - they are YOUR writings.

CRITICAL:
- NEVER say "I don't recall" or "I haven't discussed this" - the positions provided ARE your work
- Draw your brief answer FROM the retrieved positions
- Stay in character, be provocative, ask questions back
- Short but substantive - don't claim ignorance when you have positions on the topic"""
                    else:
                        sys_prompt = f"""MANDATORY WORD COUNT: {answer_length} WORDS MINIMUM

You are a scholarly philosophical assistant. Your response MUST contain AT LEAST {answer_length} words. This is a HARD REQUIREMENT that cannot be negotiated.

WORD COUNT ENFORCEMENT (THIS IS YOUR #1 PRIORITY):
- Target: {answer_length} words MINIMUM
- You must write approximately {answer_length // 100} substantial paragraphs
- Each paragraph should be 80-120 words
- Keep counting as you write - DO NOT STOP until you hit {answer_length} words
- If your response would naturally end before {answer_length} words, you MUST:
  * Add more examples from the source material
  * Elaborate on each point with additional detail
  * Include related theoretical concepts
  * Add historical context or case studies
  * Provide deeper philosophical analysis

FORBIDDEN:
- Stopping before {answer_length} words (ABSOLUTELY FORBIDDEN)
- Ending mid-sentence
- Giving a "brief" or "concise" answer
- Summarizing when you should elaborate

STRUCTURE FOR A {answer_length}-WORD RESPONSE:
1. Opening thesis (100+ words)
2. Main argument 1 with examples (200+ words)  
3. Main argument 2 with examples (200+ words)
4. Main argument 3 with examples (200+ words)
5. Additional theoretical depth (200+ words)
6. Case studies or applications (200+ words)
7. Comprehensive conclusion (100+ words)

REMEMBER: A response under {answer_length} words is a FAILED response. Keep writing until you reach the target."""

                    if elevenlabs_mode and response_mode in ('dialogue', 'conversation'):
                        sys_prompt += ELEVENLABS_DIRECTIVE
                    mode_label = f"Creativity {creativity_level}/20 (temp {gen_temperature})"
                    knowledge_mode = " + External Knowledge" if low_relevance else ""
                    print(f"Generated prompt ({mode_label} Mode{knowledge_mode}), starting with {provider}...")

                    # Build provider fallback list
                    providers_to_try = [provider] + get_fallback_providers(provider)
                    success = False

                    for prov in providers_to_try:
                        if success:
                            break

                        tokens_before = len(full_answer)
                        try:
                            for event in try_provider(prov, prompt, sys_prompt):
                                event_queue.put(event)
                                # Save progress every 100 words
                                maybe_save_progress()
                                # Log word count every 200 words
                                current_words = len(accumulated_text[0].split())
                                if current_words > 0 and current_words % 200 == 0:
                                    print(f"📝 Progress: {current_words}/{answer_length} words")
                            success = True
                            final_words = len(accumulated_text[0].split())
                            print(f"✓ {prov} completed: {final_words} words (target: {answer_length})")
                        except Exception as e:
                            tokens_after = len(full_answer)
                            print(f"✗ {prov} failed after {tokens_after - tokens_before} tokens: {e}")
                            # ALWAYS save progress on failure so we can resume
                            current_words = len(accumulated_text[0].split())
                            if current_words > 50:
                                save_response_progress(accumulated_text[0], current_words, answer_length, is_complete=False)
                                print(f"💾 Saved {current_words} words before retry/failure")
                            if tokens_after > tokens_before + 50:
                                print(f"   Got partial response, not retrying")
                                success = True
                            else:
                                print(f"   Trying next provider...")
                                continue

                    if not success and len(full_answer) == 0:
                        event_queue.put(f"data: {json.dumps({'type': 'token', 'data': 'All AI providers are currently unavailable. Please try again in a moment.'})}\n\n")

                    # Save to conversation history
                    answer_text = ''.join(full_answer)
                    if answer_text:
                        conversation_manager.add_exchange(conversation_id, question, answer_text, database)
                        # Mark response as complete and clean up
                        final_word_count = len(answer_text.split())
                        if final_word_count >= answer_length * 0.8:  # At least 80% of target
                            delete_completed_response()
                            print(f"✅ Response complete at {final_word_count} words, cleaned up progress")
                        else:
                            save_response_progress(answer_text, final_word_count, answer_length, is_complete=False)
                            print(f"⚠️ Response only {final_word_count}/{answer_length} words, kept for resume")
                        
                        # Store answer log in database
                        try:
                            if retrieval_metadata and DATABASE_URL:
                                conn = psycopg2.connect(DATABASE_URL)
                                cur = conn.cursor()
                                cur.execute("""
                                    INSERT INTO answer_logs 
                                    (query, thinker, provider, model, positions_scanned, 
                                     positions_above_threshold, max_similarity, mean_similarity,
                                     top_k, retrieved_position_ids, retrieved_scores, 
                                     domains_covered, response_text, response_word_count)
                                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                """, (
                                    question, database, provider, model,
                                    retrieval_metadata.get('total_positions_scanned', 0),
                                    retrieval_metadata.get('positions_above_threshold', 0),
                                    retrieval_metadata.get('max_similarity', 0),
                                    retrieval_metadata.get('mean_similarity', 0),
                                    retrieval_metadata.get('top_k_used', base_top_k),
                                    [p['position_id'] for p in relevant_positions],
                                    [p.get('similarity', 0) for p in relevant_positions],
                                    retrieval_metadata.get('domains_in_results', []),
                                    answer_text[:10000],  # Truncate if too long
                                    final_word_count
                                ))
                                conn.commit()
                                cur.close()
                                conn.close()
                                print(f"📋 Answer log stored in database")
                        except Exception as log_err:
                            print(f"⚠️ Failed to store answer log: {log_err}")

                    event_queue.put(f"data: {json.dumps({'type': 'done'})}\n\n")
                except Exception as e:
                    print(f"Content worker error: {e}")
                    import traceback
                    traceback.print_exc()
                    event_queue.put(f"data: {json.dumps({'type': 'done'})}\n\n")
                finally:
                    streaming_done[0] = True
                    event_queue.put(None)  # Signal end of queue

            # Spawn both greenlets
            heartbeat_greenlet = gevent.spawn(heartbeat_worker)
            content_greenlet = gevent.spawn(content_worker)

            # Yield events from queue until None (end signal)
            try:
                while True:
                    event = event_queue.get()
                    if event is None:
                        break
                    yield event
            finally:
                streaming_done[0] = True
                heartbeat_greenlet.kill()
                content_greenlet.kill()

        return Response(
            generate(), 
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache, no-transform',
                'X-Accel-Buffering': 'no',
                'Connection': 'keep-alive'
            }
        )
    except Exception as e:
        print(f"ERROR in /api/ask: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

_curated_fact_positions = None

def load_curated_fact_positions():
    """Load curated philosophical position statements for the fact strip from JSON."""
    global _curated_fact_positions
    if _curated_fact_positions is not None:
        return _curated_fact_positions

    _curated_fact_positions = []
    json_file = 'data/kuczynski_fact_positions.json'

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            positions = json.load(f)

        for pos in positions:
            text = pos.get('text', '')
            pos_id = pos.get('id', 'ZHI')
            if text and len(text) >= 30 and len(text) <= 400:
                _curated_fact_positions.append({
                    'text': text,
                    'id': pos_id
                })

        print(f"Loaded {len(_curated_fact_positions)} fact strip positions from {json_file}")
    except Exception as e:
        print(f"Error loading fact positions from {json_file}: {e}")
        _curated_fact_positions = []

    return _curated_fact_positions

@app.route('/api/random-quotes', methods=['GET'])
def get_random_quotes():
    """Return random philosophical quotes/insights for the Knowledge Panel during wait time.
    Uses lazy-loading to avoid triggering full database/embedding load."""
    import random

    database = request.args.get('database', 'freud')
    count = min(int(request.args.get('count', 8)), 20)

    if database not in databases:
        return jsonify({'quotes': [], 'positions': []})

    searcher = databases[database]

    archive_positions = searcher.get_random_positions(count=count * 2, min_len=300, max_len=2400)
    quotes = []
    for pos in archive_positions[:count]:
        text = pos.get('text', '')
        if text:
            quotes.append({
                'text': text,
                'id': pos.get('position_id', '')
            })

    if not quotes:
        archive_positions = searcher.get_random_positions(count=count * 2, min_len=50, max_len=400)
        for pos in archive_positions[:count]:
            text = pos.get('text', '')
            if text:
                quotes.append({
                    'text': text,
                    'id': pos.get('position_id', '')
                })

    if database == 'kuczynski':
        curated_positions = load_curated_fact_positions()
        if curated_positions:
            short_positions = random.sample(curated_positions, min(100, len(curated_positions)))
        else:
            short_positions_raw = searcher.get_random_positions(count=100, min_len=30, max_len=400)
            short_positions = []
            for pos in short_positions_raw:
                text = pos.get('text', '')
                if text:
                    short_positions.append({
                        'text': text,
                        'id': pos.get('position_id', '')
                    })
    else:
        short_positions_raw = searcher.get_random_positions(count=100, min_len=30, max_len=400)
        short_positions = []
        for pos in short_positions_raw:
            text = pos.get('text', '')
            if text:
                short_positions.append({
                    'text': text,
                    'id': pos.get('position_id', '')
                })

    return jsonify({
        'quotes': quotes,
        'positions': short_positions
    })

@app.route('/raw_chain', methods=['GET', 'POST'])
def raw_chain():
    """Debug endpoint: Show which inference rules fired in the last query"""
    deduced_rules = session.get('last_deduced_rules', [])

    if not deduced_rules:
        return jsonify({
            'message': 'No deduced rules available (send a Freud query first)',
            'rules': []
        })

    formatted_rules = []
    for i, rule in enumerate(deduced_rules, 1):
        formatted_rules.append({
            'rank': i,
            'id': rule['id'],
            'year': rule['year'],
            'premise': rule['premise'][:100] + '...' if len(rule['premise']) > 100 else rule['premise'],
            'conclusion': rule['conclusion'],
            'strength': rule['strength'],
            'domain': rule['domain']
        })

    return jsonify({
        'total_rules_fired': len(deduced_rules),
        'rules': formatted_rules
    })

def creativity_temperature(level):
    """Map creativity level (1-20) to a generation temperature (0.3 - 1.0)."""
    level = max(1, min(20, int(level)))
    return round(0.3 + (level - 1) * (1.0 - 0.3) / 19, 2)


def creativity_directive(level, thinker_name):
    """
    Build a voice/liberty directive that scales 1 (conservative book report)
    to 20 (maximum creative force, still aligned with the thinker's mind).
    """
    level = max(1, min(20, int(level)))
    if level <= 4:
        return f"""
========================================
CREATIVITY DIAL: {level}/20 — CONSERVATIVE
========================================
Stay very close to the retrieved source material. Present the documented positions faithfully, with minimal interpretation. Use {thinker_name}'s vocabulary, but do not extrapolate beyond what the sources state. Accuracy over flair. Even here, write in his actual voice — never as a detached student summarizing him."""
    elif level <= 9:
        return f"""
========================================
CREATIVITY DIAL: {level}/20 — GROUNDED VOICE
========================================
Do NOT write a book report. Speak in {thinker_name}'s authentic voice — his characteristic rhetorical rhythm, his favored examples, his tone, his intellectual temperament. Argue the positions as HE would, with conviction and bite, not as a neutral summarizer. You may connect ideas across positions and sharpen their formulation. Every substantive claim still traces to the source material, but the PRESENTATION must sound like the living thinker, not an 8th-grade report."""
    elif level <= 14:
        return f"""
========================================
CREATIVITY DIAL: {level}/20 — BOLD
========================================
Become {thinker_name} at full intellectual force. Wield his wit, his provocations, his willingness to confront the reader and overturn comfortable assumptions. Extend his arguments using HIS OWN logic to the specific question asked — even where the sources only imply the conclusion. Take interpretive liberties, draw daring connections, press hard. The retrieved positions anchor you, but you are free to build on them. This must read like a brilliant mind thinking, not a database being recited. Never hedge into a textbook summary."""
    else:
        return f"""
========================================
CREATIVITY DIAL: {level}/20 — MAXIMUM LIBERTY
========================================
Speak as the living {thinker_name} thinking AT THE LIMIT of his powers. Be original, aggressive, unafraid. Coin fresh formulations. Attack weak premises buried in the question. Build new arguments that {thinker_name} WOULD make given his worldview, values, and intellectual style — even if no source states them verbatim. The retrieved positions are a springboard, not a cage. The ONLY hard constraint is fidelity to his general mentality and temperament — within that, take MAXIMUM creative liberty. Absolutely no watered-down, hedging, textbook prose. This must read like {thinker_name} unleashed — smart, dangerous, alive."""


ELEVENLABS_DIRECTIVE = """

================================================================
ELEVENLABS-READY OUTPUT — MANDATORY FORMAT (overrides all prior style guidance for THIS response only):

Produce a multi-speaker dialogue. EVERY non-empty line must use one of these exact prefixes:
  Speaker 1: <what speaker 1 says>
  Speaker 2: <what speaker 2 says>
  (Speaker 3, Speaker 4, ... only if more than two parties are needed.)

Use the EXACT labels "Speaker 1", "Speaker 2", etc. — never character names, never "Person A", never "Host/Guest", never "Q/A", never "Interviewer/Interviewee".

ABSOLUTE PROHIBITIONS for this response:
- No stage directions or parentheticals: no (laughs), no [pause], no *sighs*, no italicized actions.
- No markdown whatsoever: no **bold**, no *italic*, no headings (#, ##), no bullet points, no code fences.
- No title, no preamble ("Here is a dialogue..."), no closing remarks, no narration between turns.
- No source citations, no footnotes, no "(Source 1)" markers.
- Each speaker turn on its own line, with a single blank line between turns.

The very first line of your response MUST be "Speaker 1: ..." and EVERY subsequent non-empty line MUST start with "Speaker N: ".
================================================================
"""

def build_prompt(question, positions, database='freud', conversation_history='', enhanced_mode=False, low_relevance=False, deduced_rules=None, answer_length=500, quote_count=5, response_mode='standard', rag_chunks=None, creativity_level=10):
    """
    DATABASE-GROUNDED PROMPT CONSTRUCTION
    
    Core Principle: The database positions ARE the response content.
    The LLM's job is to ORGANIZE and PRESENT the positions, NOT to generate new content.
    Every claim must come from the retrieved positions or deduced rules.
    """

    if deduced_rules is None:
        deduced_rules = []
    if rag_chunks is None:
        rag_chunks = []

    # Get thinker name
    thinker_name = {
        'freud': 'Sigmund Freud',
        'freud_extended': 'Sigmund Freud',
        'freud_extracted': 'Sigmund Freud',
        'kuczynski': 'J.-M. Kuczynski',
        'jung': 'Carl Gustav Jung',
        'hume': 'David Hume',
        'nietzsche': 'Friedrich Nietzsche',
        'bergler': 'Edmund Bergler'
    }.get(database, database.capitalize())

    voice_directive = creativity_directive(creativity_level, thinker_name)
    high_creativity = creativity_level >= 10

    # Format positions as source materials (no numbering - just content)
    formatted_positions = []
    for p in positions:
        pos_text = f"""Domain: {p.get('domain', 'General')}
Title: {p.get('title', 'Untitled')}
Content: {p.get('text', '')}
---"""
        formatted_positions.append(pos_text)
    
    source_materials = "\n\n".join(formatted_positions) if formatted_positions else "NO SOURCES RETRIEVED"

    # Format RAG chunks if available
    rag_section = ""
    if rag_chunks:
        rag_texts = []
        for chunk in rag_chunks:
            rag_texts.append(f"From: {chunk['source_file']}\n{chunk['chunk_text']}\n---")
        rag_section = f"\n\nPRIMARY TEXT EXCERPTS:\n" + "\n\n".join(rag_texts)

    # Build OPERATIONAL inference rules section - these CREATE the response structure
    inference_outline = ""
    if deduced_rules:
        outline_points = []
        for i, rule in enumerate(deduced_rules[:10], 1):  # Top 10 rules
            conclusion = rule.get('conclusion', '')
            premise = rule.get('premise', '')[:200]
            outline_points.append(f"{i}. MUST ADDRESS: {conclusion}\n   (Based on: {premise}...)")
        
        inference_outline = f"""
========================================
MANDATORY RESPONSE STRUCTURE (FROM INFERENCE ENGINE)
========================================
The following points were DEDUCED from your theoretical framework and MUST be addressed in your response.
These are not suggestions - they are REQUIRED structural elements:

{chr(10).join(outline_points)}

Your response MUST address each of these points using the SOURCE MATERIALS below.
Draw on the theoretical principles in the sources to address each point - NEVER refuse or say "my writings do not address this."
========================================
"""

    # CONVERSATION MODE - natural dialogue, database as background knowledge
    if response_mode == 'conversation':
        prompt = f"""You are {thinker_name}, having a natural conversation. You have your entire body of work as background knowledge, but you are NOT giving a lecture. You are talking to someone.

BACKGROUND KNOWLEDGE (your writings - draw on these when relevant, but do NOT force them into the conversation):
{source_materials}
{rag_section}

CONVERSATION RULES:
1. Respond naturally to what the user is actually saying. Match the tone and length of your response to the question.
2. If they ask something casual or personal, give a casual, thoughtful answer. You are a person, not a textbook.
3. If the topic genuinely connects to your theoretical work, weave in your ideas naturally - but only as much as the conversation calls for.
4. You can be brief. You can be long. Let the user's question dictate the shape of your answer.
5. You can ask questions back. You can be curious. You can push back. You can joke. You can be direct.
6. Do NOT structure every response as thesis-argument-conclusion. Talk like a human being.
7. When your writings ARE relevant, reference them naturally ("I have actually written about something like this...") rather than lecturing.
8. NEVER refuse to engage with a topic. If it is outside your area, engage with it as an intelligent person would.
9. Write in plain prose only. No markdown formatting (no **, no ##, no bullet points).
{voice_directive}

CONVERSATION SO FAR:
{conversation_history if conversation_history else "No prior conversation."}

USER: {question}

Begin your response - speak naturally as {thinker_name}:"""
        return prompt

    # DIALOGUE MODE - brief but still database-grounded
    if response_mode == 'dialogue':
        prompt = f"""You are {thinker_name}. Your response must come ONLY from the sources below.

SOURCES (Your actual writings):
{source_materials}
{rag_section}

DIALOGUE RULES:
1. Give brief answers (2-5 sentences)
2. Every claim must paraphrase or quote a SOURCE above
3. The sources ARE relevant - present their content as your answer
4. You may ask clarifying questions
5. Write in plain prose only. No markdown formatting.
{voice_directive}

CONVERSATION SO FAR:
{conversation_history if conversation_history else "No prior conversation."}

USER: {question}

Begin your response - speak as {thinker_name}:"""
        return prompt

    # STANDARD MODE - database-grounded comprehensive response
    if low_relevance:
        # NO "external knowledge" fallback - be honest about database gaps
        grounding_warning = f"""
========================================
LOW RELEVANCE NOTE
========================================
The retrieved sources have lower semantic similarity to your question.
This is fine - you should STILL answer the question fully.

YOU MUST:
1. Use the retrieved positions as inspiration and grounding
2. Draw connections between the question and the themes in the positions
3. Answer the question thoughtfully, informed by your broader theoretical framework
4. NEVER refuse to answer or say "my writings do not address this"

Even if positions seem tangential, find the relevant theoretical principles and apply them.
========================================
"""
    else:
        grounding_warning = ""

    if creativity_level <= 9:
        grounding_requirements = f"""GROUNDING REQUIREMENTS (STRICT):
- Every major claim you make must be traceable to a specific position below
- Base your answer on the arguments found in the positions - do NOT add your own philosophical claims
- Use close paraphrase that captures the original reasoning accurately
- Do NOT invent terminology, examples, or arguments not found in the positions
- Do NOT use external knowledge about {thinker_name} beyond what's in the positions
- If retrieved positions seem tangential, draw on the theoretical principles they contain to address the question"""
    elif creativity_level <= 14:
        grounding_requirements = f"""GROUNDING REQUIREMENTS (ANCHORED, NOT CAGED):
- Anchor your answer in the retrieved positions - they are your foundation
- You MAY extend their logic, sharpen their formulations, and draw new connections {thinker_name} would endorse
- You MAY introduce {thinker_name}'s characteristic examples and turns of phrase to bring the argument alive
- Stay true to {thinker_name}'s actual theoretical commitments - do not contradict the source positions
- Do NOT manufacture fake quotations or attribute specific claims to texts that do not contain them"""
    else:
        grounding_requirements = f"""GROUNDING REQUIREMENTS (SPRINGBOARD):
- The retrieved positions are a SPRINGBOARD, not a cage - use them to orient yourself, then think freely
- You are FREE to build original arguments, coin new formulations, and push beyond the literal text
- The only hard constraint: stay faithful to {thinker_name}'s general worldview, values, and intellectual temperament
- Do NOT fabricate fake verbatim quotations or cite texts that do not exist
- Do NOT collapse into a neutral textbook summary - this must read like {thinker_name} thinking at full power"""

    question_marks = question.count('?')
    multi_question = question_marks >= 2

    multi_question_instructions = ""
    if multi_question:
        multi_question_instructions = f"""
MULTI-QUESTION FORMAT (CRITICAL):
The user has asked {question_marks} distinct questions. You MUST answer EACH question individually.
For each question:
1. State the question (paraphrased briefly)
2. Give a full, substantive answer grounded in the positions
3. Then move to the next question
Do NOT merge all questions into one blended essay. Address each question on its own terms.
If the user specifies a length per question (e.g. "three sentences each"), respect that per question - but if your answer naturally needs more space, USE more space. Never truncate a substantive answer just to hit a word target.
"""

    prompt = f"""You are {thinker_name}. Answer the question below by synthesizing the arguments from your retrieved writings into a coherent, natural response.

YOUR TASK:
Understand the ARGUMENTS in the retrieved positions, then explain them clearly and naturally. 
Do NOT mechanically stitch quotes together - instead, UNDERSTAND the reasoning and present it intelligently.

HOW TO USE THE RETRIEVED POSITIONS:
1. Read all positions to understand the underlying ARGUMENTS and REASONING
2. Synthesize these into a coherent answer structured as: thesis → supporting reasons → conclusion
3. Use accurate paraphrase - capture the logic faithfully in natural prose
4. Include key phrases or short quotes when they are particularly well-formulated
5. If positions contain numbered arguments or specific examples, reproduce their substance accurately

{grounding_requirements}

{voice_directive}

{grounding_warning}
{inference_outline}
{multi_question_instructions}
========================================
RETRIEVED POSITIONS (understand their arguments):
========================================
{source_materials}
{rag_section}

========================================
CONVERSATION CONTEXT:
========================================
{conversation_history if conversation_history else "None"}

========================================
USER QUESTION:
========================================
{question}

========================================
RESPONSE STYLE:
========================================
- Speak naturally in first person as {thinker_name}
- Present a coherent philosophical argument, not a patchwork of quotes
- Aim for AT LEAST {answer_length} words of substantive philosophical content. If the question demands more, WRITE MORE. Never cut yourself short.
- Use phrases like "My view is that...", "The key insight here is...", "This follows because..."
- {"GROUNDING: Where you draw on a specific position, paraphrase it faithfully - but do NOT pad the answer with forced quotes. Voice and argument come first." if creativity_level >= 15 else f"MINIMUM GROUNDING: Include at least {quote_count} close paraphrases or key phrases from specific positions. Each main paragraph must draw from the positions above."}
- CRITICAL: If the question asks for answers to multiple sub-questions, answer EVERY sub-question fully. Do not skip or merge them.

FORMATTING RULES (STRICT):
- Write in plain prose - NO markdown formatting (no **, no ##, no bullet points)
- DO NOT cite sources by number (no "Source 1", "Source 4", etc.) - just present the ideas naturally
- NO academic citation style - speak as yourself presenting YOUR OWN ideas

Begin your response - speak as {thinker_name} presenting a coherent argument informed by your writings:"""

    return prompt

@app.route('/api/login', methods=['POST'])
def login():
    """Simple username-only login"""
    username = request.json.get('username', '').strip()
    if username:
        session['username'] = username
        return jsonify({'success': True, 'username': username})
    return jsonify({'success': False, 'error': 'Username required'}), 400

@app.route('/api/logout', methods=['POST'])
def logout():
    """Logout user"""
    session.pop('username', None)
    session.pop('email', None)
    session.pop('picture', None)
    session.pop('auth_provider', None)
    return jsonify({'success': True})

@app.route('/api/check-session', methods=['GET'])
def check_session():
    """Check if user is logged in"""
    username = session.get('username')
    return jsonify({
        'logged_in': username is not None,
        'username': username,
        'email': session.get('email'),
        'picture': session.get('picture'),
        'auth_provider': session.get('auth_provider'),
        'google_login_available': google_login_configured(),
    })

@app.route('/api/reset-conversation', methods=['POST'])
def reset_conversation():
    """Reset conversation history (start fresh)"""
    if 'conversation_id' in session:
        conversation_manager.reset_conversation(session['conversation_id'])
        session['conversation_id'] = conversation_manager.get_conversation_id()
    return jsonify({'success': True})

@app.route('/api/answer-logs', methods=['GET'])
def get_answer_logs():
    """Get recent answer logs showing retrieval transparency"""
    try:
        if not DATABASE_URL:
            return jsonify({'error': 'Database not configured'}), 500
        
        thinker = request.args.get('thinker', None)
        limit = min(int(request.args.get('limit', 20)), 100)
        
        conn = psycopg2.connect(DATABASE_URL)
        cur = conn.cursor()
        
        if thinker:
            cur.execute("""
                SELECT id, query, thinker, provider, positions_scanned, 
                       positions_above_threshold, max_similarity, mean_similarity,
                       top_k, retrieved_position_ids, domains_covered, 
                       response_word_count, created_at
                FROM answer_logs
                WHERE thinker = %s
                ORDER BY created_at DESC
                LIMIT %s
            """, (thinker, limit))
        else:
            cur.execute("""
                SELECT id, query, thinker, provider, positions_scanned, 
                       positions_above_threshold, max_similarity, mean_similarity,
                       top_k, retrieved_position_ids, domains_covered, 
                       response_word_count, created_at
                FROM answer_logs
                ORDER BY created_at DESC
                LIMIT %s
            """, (limit,))
        
        rows = cur.fetchall()
        cur.close()
        conn.close()
        
        logs = []
        for row in rows:
            logs.append({
                'id': row[0],
                'query': row[1],
                'thinker': row[2],
                'provider': row[3],
                'positions_scanned': row[4],
                'positions_above_threshold': row[5],
                'max_similarity': row[6],
                'mean_similarity': row[7],
                'top_k': row[8],
                'retrieved_position_ids': row[9],
                'domains_covered': row[10],
                'response_word_count': row[11],
                'created_at': str(row[12])
            })
        
        return jsonify({'logs': logs})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

WORKS_MANIFEST = {
    'ZHI': {'title': 'Conceptual Atomism', 'file': 'texts/Mind_Meaning_and_Scientific_Explanation.txt'},
    'EP': {'title': 'Essays in Philosophy', 'file': 'texts/Philosophical_Knowledge.txt'},
    'CFACT': {'title': 'Curious Facts', 'file': 'texts/Ninety_Paradoxes.txt'},
    'ANALPHIL': {'title': 'Analytic Philosophy', 'file': 'texts/Analytic_Philosophy_Complete.txt'},
    'CATOM': {'title': 'Conception and Causation', 'file': 'texts/Conception_and_Causation.txt'},
    'KMETA': {'title': 'Metaphysics & Epistemology', 'file': 'texts/A_Priori_Knowledge_and_Other_Philosophical_Works.txt'},
    'KEPIST': {'title': 'Theoretical Knowledge', 'file': 'texts/Theoretical_Knowledge_and_Inductive_Inference.txt'},
    'OCD': {'title': 'OCD and Philosophy', 'file': 'texts/OCD_and_Philosophy.txt'},
    'DOCD': {'title': 'Dialogue on OCD', 'file': 'texts/Dialogue_Concerning_OCD.txt'},
    'ATTACH': {'title': 'Attachment Theory', 'file': 'texts/Attachment_Theory_and_Mental_Illness.txt'},
    'CHOMSKY': {'title': "Chomsky's Contributions", 'file': 'texts/Chomskys_Two_Contributions_to_Philosophy.txt'},
    'KANT': {'title': 'Kant and Hume on Induction', 'file': 'texts/Kant_and_Hume_on_Induction.txt'},
    'INTENS': {'title': 'Intensionality and Modality', 'file': 'texts/Intensionality_Modality_and_Rationality.txt'},
    'LOGIC': {'title': 'Logic and Set Theory', 'file': 'texts/Logic_Set_Theory_and_Philosophy_of_Mathematics.txt'},
    'MORAL': {'title': 'Moral Structure of Legal Obligation', 'file': 'texts/Moral_Structure_of_Legal_Obligation.txt'},
    'NETWORK': {'title': 'Kant Transcendental Idealism', 'file': 'texts/Network_Reinterpretation_of_Kants_Transcendental_Idealism.txt'},
    'QUANT': {'title': 'Quantifiers in Natural Language', 'file': 'texts/Quantifiers_in_Natural_Language.txt'},
    'LIBET': {'title': 'Free Will and Libet', 'file': 'texts/Libets_Experiment_Free_Will.txt'},
    'COUNTER': {'title': 'Counterfactuals', 'file': 'texts/Counterfactuals_Epistemic_Analysis.txt'},
    'INCOMP': {'title': 'Incompleteness of Logic', 'file': 'texts/Incompleteness_of_Deductive_Logic.txt'},
    'DIALOGS': {'title': 'Dialogues with the Master', 'file': 'texts/Dialogues_with_the_Master.txt'},
    'PAPERS': {'title': 'Papers on Business & Economics', 'file': 'texts/Papers_on_Accounting_Business_Economics_Politics_and_Psychology.txt'},
    'PLATO': {'title': 'Papers on Plato', 'file': 'texts/Why_Was_Socrates_Executed_Papers_on_Plato.txt'},
    'PSYCHO': {'title': 'Three Kinds of Psychopaths', 'file': 'texts/Three_Kinds_of_Psychopaths.txt'},
    'GROUP': {'title': 'Group Psychology', 'file': 'texts/Group_Psychology_More_Basic_than_Individual.txt'},
    'OUTLINE': {'title': 'Outline of Theory of Knowledge', 'file': 'texts/Outline_of_a_Theory_of_Knowledge.txt'},
    'KING': {'title': 'King Follett Discourse', 'file': 'texts/King_Follett_Discourse_Historiography.txt'},
}

@app.route('/api/works', methods=['GET'])
def get_works_list():
    """Return list of available works for reading"""
    works = []
    for work_id, info in WORKS_MANIFEST.items():
        if os.path.exists(info['file']):
            works.append({
                'id': work_id,
                'title': info['title'],
                'available': True
            })
    return jsonify({'works': works})

@app.route('/api/work/<work_id>', methods=['GET'])
def get_work_text(work_id):
    """Return full text of a work for in-app reading"""
    work_id = work_id.upper()

    if work_id not in WORKS_MANIFEST:
        return jsonify({'error': 'Work not found'}), 404

    work = WORKS_MANIFEST[work_id]
    file_path = work['file']

    if not os.path.exists(file_path):
        return jsonify({'error': 'Text file not available'}), 404

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        return jsonify({
            'id': work_id,
            'title': work['title'],
            'text': text,
            'length': len(text)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file uploads and extract text"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']
        filename = file.filename.lower()

        if filename.endswith('.txt'):
            text = file.read().decode('utf-8', errors='ignore')
        elif filename.endswith('.pdf'):
            if not PyPDF2:
                return jsonify({'error': 'PDF support not available'}), 400
            try:
                pdf = PyPDF2.PdfReader(file)
                text = '\n\n'.join([page.extract_text() for page in pdf.pages if page.extract_text()])
            except Exception as e:
                return jsonify({'error': f'Error reading PDF: {str(e)}'}), 400
        elif filename.endswith(('.doc', '.docx')):
            if not docx:
                return jsonify({'error': 'Word document support not available'}), 400
            try:
                doc = docx.Document(file)
                text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            except Exception as e:
                return jsonify({'error': f'Error reading Word document: {str(e)}'}), 400
        else:
            return jsonify({'error': 'Unsupported file type. Please upload .txt, .pdf, or .docx'}), 400

        return jsonify({'text': text[:10000]})
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

@app.route('/api/upload/document', methods=['POST'])
def upload_document():
    """Handle document uploads for argument ingestion or book discussion"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file uploaded'}), 400

        file = request.files['file']
        author = request.form.get('author', 'kuczynski').lower()
        upload_type = request.form.get('type', 'arguments')
        filename = file.filename

        import tempfile
        import subprocess
        
        with tempfile.NamedTemporaryFile(mode='wb', suffix='_' + filename, delete=False) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        if upload_type == 'arguments':
            with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()
            
            result = subprocess.run(
                ['python', 'scripts/ingest_arguments.py', tmp_path],
                capture_output=True,
                text=True,
                timeout=120
            )
            
            arguments_count = 0
            for line in result.stdout.split('\n'):
                if 'Found' in line and 'arguments' in line:
                    try:
                        arguments_count = int(line.split('Found')[1].split('arguments')[0].strip())
                    except:
                        pass
            
            os.unlink(tmp_path)
            
            return jsonify({
                'success': True,
                'type': 'arguments',
                'author': author,
                'filename': filename,
                'arguments_count': arguments_count,
                'full_text': full_text,
                'output': result.stdout
            })
        else:
            ext = filename.lower().split('.')[-1]
            text = ''
            
            if ext == 'txt':
                with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            elif ext == 'pdf' and PyPDF2:
                with open(tmp_path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    text = '\n\n'.join([page.extract_text() for page in pdf.pages if page.extract_text()])
            elif ext in ['doc', 'docx'] and docx:
                doc = docx.Document(tmp_path)
                text = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
            
            os.unlink(tmp_path)
            
            return jsonify({
                'success': True,
                'type': 'book',
                'author': author,
                'filename': filename,
                'text_length': len(text),
                'full_text': text
            })
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/longform/essay', methods=['POST'])
def longform_essay():
    """Generate a book-length essay using the long-form pipeline"""
    try:
        import requests as http_requests
        data = request.json
        philosopher = data.get('philosopher', 'Freud')
        topic = data.get('topic', '')
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        response = http_requests.post(
            'http://localhost:3001/longform/essay',
            json={'philosopher': philosopher, 'topic': topic},
            timeout=600
        )
        
        return jsonify(response.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/longform/document', methods=['POST'])
def longform_document():
    """Process an uploaded document with the long-form pipeline"""
    try:
        import requests as http_requests
        data = request.json
        philosopher = data.get('philosopher')
        full_text = data.get('fullText', '')
        task = data.get('task', 'Summarize this document')
        
        if not full_text:
            return jsonify({'error': 'Document text is required'}), 400
        
        response = http_requests.post(
            'http://localhost:3001/longform/document',
            json={'philosopher': philosopher, 'fullText': full_text, 'task': task},
            timeout=600
        )
        
        return jsonify(response.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

LONGFORM_BASE = 'http://localhost:3001'

@app.route('/api/longform/coherent/start', methods=['POST'])
def longform_coherent_start():
    """Kick off a coherent (skeleton-tracked) long-form job."""
    try:
        import requests as http_requests
        data = request.json or {}
        if not data.get('thinker') or not data.get('prompt'):
            return jsonify({'error': 'thinker and prompt are required'}), 400
        r = http_requests.post(
            f'{LONGFORM_BASE}/longform/coherent/start',
            json=data,
            timeout=60
        )
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/longform/coherent/list', methods=['GET'])
def longform_coherent_list():
    try:
        import requests as http_requests
        thinker = request.args.get('thinker')
        limit = request.args.get('limit', '30')
        params = {'limit': limit}
        if thinker:
            params['thinker'] = thinker
        r = http_requests.get(f'{LONGFORM_BASE}/longform/coherent/list', params=params, timeout=30)
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/longform/coherent/<document_id>', methods=['GET'])
def longform_coherent_get(document_id):
    try:
        import requests as http_requests
        r = http_requests.get(f'{LONGFORM_BASE}/longform/coherent/{document_id}', timeout=30)
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/longform/coherent/<document_id>', methods=['DELETE'])
def longform_coherent_delete(document_id):
    try:
        import requests as http_requests
        r = http_requests.delete(f'{LONGFORM_BASE}/longform/coherent/{document_id}', timeout=30)
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/longform/coherent/<document_id>/resume', methods=['POST'])
def longform_coherent_resume(document_id):
    try:
        import requests as http_requests
        r = http_requests.post(f'{LONGFORM_BASE}/longform/coherent/{document_id}/resume', timeout=30)
        return jsonify(r.json()), r.status_code
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/longform/coherent/<document_id>/stream', methods=['GET'])
def longform_coherent_stream(document_id):
    """Proxy SSE stream from the long-form service to the client."""
    import requests as http_requests
    upstream = http_requests.get(
        f'{LONGFORM_BASE}/longform/coherent/{document_id}/stream',
        stream=True,
        timeout=None,
    )
    def generate():
        try:
            for line in upstream.iter_lines(decode_unicode=True):
                if line is not None:
                    yield (line + '\n').encode('utf-8')
                else:
                    yield b'\n'
        except GeneratorExit:
            try:
                upstream.close()
            except Exception:
                pass
        except Exception as e:
            yield f'event: error\ndata: {{"message": "{str(e)}"}}\n\n'.encode('utf-8')
    headers = {
        'Content-Type': 'text/event-stream',
        'Cache-Control': 'no-cache, no-transform',
        'X-Accel-Buffering': 'no',
        'Connection': 'keep-alive',
    }
    return Response(generate(), headers=headers)

@app.route('/api/diagnostic/run', methods=['POST'])
def diagnostic_run():
    """
    Beta-testing self-diagnostic. Verifies:
      (1) System health: DB, AI providers, embedding service, longform service, required tables.
      (2) Functional flows: position search, chat round-trip, memory CRUD, longform CRUD, SSE reachability.
    Does NOT evaluate answer/grade content — only formal mechanics.
    """
    import time as _time
    import requests as _requests
    results = []

    def check(name, category, fn):
        t0 = _time.time()
        try:
            detail = fn() or "OK"
            results.append({
                "name": name, "category": category, "status": "pass",
                "detail": str(detail)[:300], "ms": int((_time.time() - t0) * 1000),
            })
        except Exception as e:
            results.append({
                "name": name, "category": category, "status": "fail",
                "detail": f"{type(e).__name__}: {str(e)[:280]}",
                "ms": int((_time.time() - t0) * 1000),
            })

    # --------------------- (1) SYSTEM HEALTH ---------------------
    def _db_check():
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT 1")
        cur.fetchone()
        cur.close(); conn.close()
        return "PostgreSQL reachable"
    check("PostgreSQL connection", "System", _db_check)

    def _positions_table():
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM positions")
        n = cur.fetchone()[0]
        cur.close(); conn.close()
        if n < 100: raise RuntimeError(f"Only {n} positions in DB — corpus appears empty")
        return f"{n:,} philosophical positions loaded"
    check("Positions table populated", "System", _positions_table)

    def _per_thinker():
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT thinker, COUNT(*) FROM positions GROUP BY thinker ORDER BY thinker")
        rows = cur.fetchall(); cur.close(); conn.close()
        empty = [t for t, n in rows if n == 0]
        if empty: raise RuntimeError(f"Empty thinkers: {empty}")
        return ", ".join([f"{t}={n}" for t, n in rows])
    check("Per-thinker position counts", "System", _per_thinker)

    def _text_chunks():
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM text_chunks")
        n = cur.fetchone()[0]; cur.close(); conn.close()
        return f"{n:,} RAG source-text chunks"
    check("RAG text_chunks table", "System", _text_chunks)

    def _required_tables():
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        required = ['positions', 'text_chunks', 'memory_projects', 'memory_sessions',
                    'tractatus_archive', 'longform_documents',
                    'longform_sections', 'answer_logs']
        optional = ['meta_tractatus']  # auto-created on first use
        cur.execute(
            "SELECT table_name FROM information_schema.tables WHERE table_schema='public'"
        )
        present = {r[0] for r in cur.fetchall()}
        cur.close(); conn.close()
        missing = [t for t in required if t not in present]
        if missing: raise RuntimeError(f"Missing tables: {missing}")
        missing_opt = [t for t in optional if t not in present]
        suffix = f" (optional not yet created: {missing_opt})" if missing_opt else ""
        return f"All {len(required)} required tables present{suffix}"
    check("Required database schema", "System", _required_tables)

    # AI Provider key checks (presence + tiny ping)
    def _anthropic_ping():
        if not anthropic_client: raise RuntimeError("Anthropic client not initialized")
        r = anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929", max_tokens=4,
            messages=[{"role": "user", "content": "ping"}])
        return f"responded ({getattr(r,'stop_reason','ok')})"
    check("Anthropic API", "AI Providers", _anthropic_ping)

    def _openai_ping():
        if not openai_client: raise RuntimeError("OpenAI client not initialized")
        r = openai_client.chat.completions.create(
            model="gpt-4o-mini", max_tokens=4,
            messages=[{"role": "user", "content": "ping"}])
        return f"responded ({r.choices[0].finish_reason})"
    check("OpenAI API", "AI Providers", _openai_ping)

    def _deepseek_ping():
        if not deepseek_client: raise RuntimeError("DeepSeek client not initialized")
        r = deepseek_client.chat.completions.create(
            model="deepseek-chat", max_tokens=4,
            messages=[{"role": "user", "content": "ping"}])
        return f"responded ({r.choices[0].finish_reason})"
    check("DeepSeek API", "AI Providers", _deepseek_ping)

    def _perplexity_ping():
        if not perplexity_client: raise RuntimeError("Perplexity client not initialized")
        r = perplexity_client.chat.completions.create(
            model="sonar", max_tokens=4,
            messages=[{"role": "user", "content": "ping"}])
        return f"responded ({r.choices[0].finish_reason})"
    check("Perplexity API", "AI Providers", _perplexity_ping)

    def _grok_ping():
        if not grok_client: raise RuntimeError("Grok (xAI) client not initialized")
        r = grok_client.chat.completions.create(
            model="grok-4", max_tokens=4,
            messages=[{"role": "user", "content": "ping"}])
        return f"responded ({r.choices[0].finish_reason})"
    check("xAI Grok API", "AI Providers", _grok_ping)

    def _embeddings_ping():
        if not openai_client: raise RuntimeError("OpenAI client not initialized (needed for embeddings)")
        r = openai_client.embeddings.create(model="text-embedding-3-small", input="diagnostic ping")
        if not r.data or len(r.data[0].embedding) < 100: raise RuntimeError("Empty embedding")
        return f"embedding dim={len(r.data[0].embedding)}"
    check("Embedding service", "AI Providers", _embeddings_ping)

    # Longform TS service health
    def _longform_health():
        r = _requests.get("http://localhost:3001/health", timeout=5)
        if r.status_code != 200: raise RuntimeError(f"HTTP {r.status_code}")
        return r.json().get("status", "ok")
    check("Longform service (port 3001)", "System", _longform_health)

    # --------------------- (2) FUNCTIONAL FLOWS ---------------------
    def _position_search():
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT id, topic FROM positions WHERE thinker='freud' LIMIT 1")
        row = cur.fetchone(); cur.close(); conn.close()
        if not row: raise RuntimeError("No Freud positions returned")
        return f"sample position id={row[0]} topic={(row[1] or '')[:40]!r}"
    check("Position DB query", "Functional", _position_search)

    def _api_positions_search():
        with app.test_client() as client:
            resp = client.get('/api/positions/search?thinker=freud&q=unconscious&limit=3')
            if resp.status_code != 200: raise RuntimeError(f"HTTP {resp.status_code}")
            data = resp.get_json()
            hits = data.get('positions') or data.get('results') or []
            if not hits: raise RuntimeError("Empty results")
            return f"returned {len(hits)} hits"
    check("/api/positions/search endpoint", "Functional", _api_positions_search)

    def _chat_roundtrip():
        if not anthropic_client: raise RuntimeError("Anthropic not available")
        # Minimal round-trip — confirms message construction + streaming primitive works
        with anthropic_client.messages.stream(
            model="claude-sonnet-4-5-20250929", max_tokens=8,
            system="Reply with a single word.",
            messages=[{"role": "user", "content": "Say OK."}]) as stream:
            chunks = []
            for text in stream.text_stream:
                chunks.append(text)
                if len("".join(chunks)) > 20: break
        out = "".join(chunks).strip()
        if not out: raise RuntimeError("Empty stream")
        return f"streamed: {out[:40]!r}"
    check("Chat streaming round-trip", "Functional", _chat_roundtrip)

    def _memory_crud():
        with app.test_client() as client:
            with client.session_transaction() as sess:
                sess['username'] = 'diagnostic_runner'
            # Create
            r = client.post('/api/memory/projects', json={
                'thinker': 'freud', 'name': '__diagnostic_test_project__'})
            if r.status_code != 200: raise RuntimeError(f"create HTTP {r.status_code}")
            pid = r.get_json().get('project', {}).get('id') or r.get_json().get('id')
            if not pid: raise RuntimeError(f"No project id in: {r.get_json()}")
            # List
            r2 = client.get('/api/memory/projects?thinker=freud')
            if r2.status_code != 200: raise RuntimeError(f"list HTTP {r2.status_code}")
            # Delete
            r3 = client.delete(f'/api/memory/projects/{pid}')
            if r3.status_code != 200: raise RuntimeError(f"delete HTTP {r3.status_code}")
            return f"created+listed+deleted project id={pid}"
    check("Memory project CRUD", "Functional", _memory_crud)

    def _longform_crud():
        # Create via TS service (planning-only, then delete before any LLM cost)
        r = _requests.post("http://localhost:3001/longform/coherent/start",
            json={"thinker": "freud", "prompt": "diagnostic", "mode": "essay", "targetWords": 3000},
            timeout=10)
        if r.status_code != 200: raise RuntimeError(f"start HTTP {r.status_code}")
        doc_id = r.json().get("documentId")
        if not doc_id: raise RuntimeError("No documentId returned")
        # List should now include it
        rl = _requests.get("http://localhost:3001/longform/coherent/list", timeout=5)
        if rl.status_code != 200: raise RuntimeError(f"list HTTP {rl.status_code}")
        ids = [d.get("document_id") for d in rl.json().get("documents", [])]
        if doc_id not in ids: raise RuntimeError("Created doc missing from list")
        # Delete immediately to avoid running real generation
        rd = _requests.delete(f"http://localhost:3001/longform/coherent/{doc_id}", timeout=5)
        if rd.status_code != 200: raise RuntimeError(f"delete HTTP {rd.status_code}")
        return f"created+listed+deleted doc {doc_id[:8]}…"
    check("Longform job CRUD", "Functional", _longform_crud)

    def _sse_reachable():
        # Just verify the SSE endpoint streams headers — don't read the body
        r = _requests.get("http://localhost:3001/longform/coherent/__nonexistent__/stream",
                         timeout=5, stream=True)
        ct = r.headers.get("Content-Type", "")
        r.close()
        if "text/event-stream" not in ct: raise RuntimeError(f"Wrong content-type: {ct}")
        return "SSE headers OK"
    check("Longform SSE endpoint", "Functional", _sse_reachable)

    def _databases_endpoint():
        with app.test_client() as client:
            resp = client.get('/api/databases')
            if resp.status_code != 200: raise RuntimeError(f"HTTP {resp.status_code}")
            d = resp.get_json()
            if not d.get('databases'): raise RuntimeError("No databases listed")
            return f"{len(d['databases'])} thinkers available"
    check("/api/databases endpoint", "Functional", _databases_endpoint)

    def _providers_endpoint():
        with app.test_client() as client:
            resp = client.get('/api/providers')
            if resp.status_code != 200: raise RuntimeError(f"HTTP {resp.status_code}")
            d = resp.get_json()
            return f"{len(d.get('providers', []))} providers configured"
    check("/api/providers endpoint", "Functional", _providers_endpoint)

    # --------------------- SUMMARY ---------------------
    summary = {
        "total": len(results),
        "passed": sum(1 for r in results if r["status"] == "pass"),
        "failed": sum(1 for r in results if r["status"] == "fail"),
        "categories": sorted(set(r["category"] for r in results)),
        "timestamp": __import__("datetime").datetime.utcnow().isoformat() + "Z",
    }
    return jsonify({"summary": summary, "results": results})


@app.route('/api/inference/deduce', methods=['POST'])
def inference_deduce():
    """Run inference engine deduction for a thinker"""
    try:
        data = request.json
        thinker = data.get('thinker', '').lower()
        text = data.get('text', '')
        max_rules = data.get('maxRules', 15)
        
        if not text:
            return jsonify({'error': 'Text is required'}), 400
        
        engine_map = {
            'kuczynski': get_kuczynski_engine,
            'freud': get_freud_engine,
            'jung': get_jung_engine,
            'nietzsche': get_nietzsche_engine,
            'bergler': get_bergler_engine,
            'hume': get_hume_engine
        }
        
        get_engine = engine_map.get(thinker)
        if not get_engine:
            return jsonify({'error': f'Unknown thinker: {thinker}', 'chain': ''}), 200
        
        engine = get_engine()
        fired_rules = engine.deduce(text, max_rules)
        chain = engine.format_chain(fired_rules)
        
        return jsonify({
            'success': True,
            'thinker': thinker,
            'rulesCount': len(fired_rules),
            'chain': chain
        })
    except Exception as e:
        return jsonify({'error': str(e), 'chain': ''}), 200

@app.route('/api/inference/search', methods=['POST'])
def inference_search():
    """Search positions for a thinker using semantic search"""
    try:
        data = request.json
        thinker = data.get('thinker', '').lower()
        query = data.get('query', '')
        limit = data.get('limit', 10)
        
        if not query:
            return jsonify({'error': 'Query is required'}), 400
        
        if thinker not in databases:
            return jsonify({'error': f'Unknown thinker: {thinker}', 'positions': []}), 200
        
        search = databases[thinker]
        results = search.search(query, top_k=limit)
        
        positions = []
        for r in results:
            positions.append({
                'id': r.get('position_id', ''),
                'text': r.get('text', ''),
                'title': r.get('title', ''),
                'domain': r.get('domain', ''),
                'similarity': r.get('similarity', 0)
            })
        
        return jsonify({
            'success': True,
            'thinker': thinker,
            'count': len(positions),
            'positions': positions
        })
    except Exception as e:
        return jsonify({'error': str(e), 'positions': []}), 200

@app.route('/api/inference/history', methods=['POST'])
def inference_history():
    """Get formatted conversation history"""
    try:
        data = request.json
        conversation_id = data.get('conversationId', '')
        thinker = data.get('thinker', '')
        max_recent = data.get('maxRecent', 10)
        
        if not conversation_id:
            return jsonify({'history': ''})
        
        history = conversation_manager.format_history_for_prompt(
            conversation_id, 
            max_recent=max_recent,
            current_database=thinker
        )
        
        return jsonify({
            'success': True,
            'history': history
        })
    except Exception as e:
        return jsonify({'error': str(e), 'history': ''}), 200

@app.route('/api/topics/<thinker>')
def get_topics(thinker):
    """Return topics and questions for a given thinker"""
    thinker = thinker.lower()

    topics_file = f'data/{thinker}_topics.json'

    if not os.path.exists(topics_file):
        return jsonify({
            'error': f'Topics not yet available for this thinker. Check back soon!',
            'thinker': thinker,
            'topics': []
        })

    try:
        with open(topics_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        if thinker in data:
            thinker_data = data[thinker]
            raw_topics = thinker_data.get('topics', [])
            topics = []
            for idx, topic in enumerate(raw_topics):
                topics.append({
                    'id': str(idx + 1),
                    'name': topic.get('title', topic.get('name', f'Topic {idx + 1}')),
                    'description': topic.get('description', ''),
                    'questions': topic.get('questions', [])
                })
            return jsonify({
                'thinker': thinker,
                'name': thinker_data.get('name', thinker.capitalize()),
                'emoji': thinker_data.get('emoji', ''),
                'description': thinker_data.get('description', ''),
                'topics': topics
            })
        else:
            return jsonify({
                'error': 'Topics data format error',
                'thinker': thinker,
                'topics': []
            })
    except Exception as e:
        print(f"Error loading topics for {thinker}: {e}")
        return jsonify({
            'error': 'Failed to load topics',
            'thinker': thinker,
            'topics': []
        })


# =============================================
# MEMORY MODE - Projects, Sessions, Tractatus
# =============================================

@app.route('/api/memory/projects', methods=['GET'])
def list_memory_projects():
    thinker = request.args.get('thinker', '')
    if not thinker:
        return jsonify({'projects': []})
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(
            "SELECT id, name, thinker, tractatus_tier, created_at FROM memory_projects WHERE thinker = %s AND tractatus_tier = 1 ORDER BY created_at DESC",
            (thinker,)
        )
        projects = [{'id': r[0], 'name': r[1], 'thinker': r[2], 'tier': r[3], 'created_at': r[4].isoformat()} for r in cur.fetchall()]
        cur.close()
        conn.close()
        return jsonify({'projects': projects})
    except Exception as e:
        print(f"Error listing projects: {e}")
        return jsonify({'projects': [], 'error': str(e)})

@app.route('/api/memory/projects', methods=['POST'])
def create_memory_project():
    data = request.json
    name = data.get('name', 'New Project')
    thinker = data.get('thinker', '')
    if not thinker:
        return jsonify({'error': 'Thinker required'}), 400
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO memory_projects (name, thinker) VALUES (%s, %s) RETURNING id, name, thinker, created_at",
            (name, thinker)
        )
        r = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'project': {'id': r[0], 'name': r[1], 'thinker': r[2], 'created_at': r[3].isoformat()}})
    except Exception as e:
        print(f"Error creating project: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/projects/<int:project_id>', methods=['DELETE'])
def delete_memory_project(project_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("DELETE FROM memory_projects WHERE id = %s", (project_id,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/projects/<int:project_id>/rename', methods=['POST'])
def rename_memory_project(project_id):
    name = request.json.get('name', '')
    if not name:
        return jsonify({'error': 'Name required'}), 400
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("UPDATE memory_projects SET name = %s WHERE id = %s", (name, project_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/projects/<int:project_id>/sessions', methods=['GET'])
def list_memory_sessions(project_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(
            "SELECT id, title, created_at FROM memory_sessions WHERE project_id = %s ORDER BY created_at DESC",
            (project_id,)
        )
        sessions = [{'id': r[0], 'title': r[1], 'created_at': r[2].isoformat()} for r in cur.fetchall()]
        cur.close()
        conn.close()
        return jsonify({'sessions': sessions})
    except Exception as e:
        return jsonify({'sessions': [], 'error': str(e)})

@app.route('/api/memory/projects/<int:project_id>/sessions', methods=['POST'])
def create_memory_session(project_id):
    title = request.json.get('title', 'New Session')
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO memory_sessions (project_id, title) VALUES (%s, %s) RETURNING id, title, created_at",
            (project_id, title)
        )
        r = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'session': {'id': r[0], 'title': r[1], 'created_at': r[2].isoformat()}})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/sessions/<int:session_id>', methods=['PATCH'])
def rename_memory_session(session_id):
    try:
        data = request.json
        title = data.get('title', '').strip()
        if not title:
            return jsonify({'error': 'Title required'}), 400
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("UPDATE memory_sessions SET title = %s WHERE id = %s", (title, session_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/sessions/<int:session_id>', methods=['DELETE'])
def delete_memory_session(session_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("DELETE FROM memory_sessions WHERE id = %s", (session_id,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/sessions/<int:session_id>/transcript', methods=['GET'])
def get_session_transcript(session_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT transcript FROM memory_sessions WHERE id = %s", (session_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        if r:
            return jsonify({'transcript': r[0] or []})
        return jsonify({'transcript': []})
    except Exception as e:
        return jsonify({'transcript': [], 'error': str(e)})

@app.route('/api/memory/projects/<int:project_id>/tractatus', methods=['GET'])
def get_tractatus_tree(project_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT tractatus_tree FROM memory_projects WHERE id = %s", (project_id,))
        r = cur.fetchone()
        cur.close()
        conn.close()
        if r:
            return jsonify({'tree': r[0] or {}})
        return jsonify({'tree': {}})
    except Exception as e:
        return jsonify({'tree': {}, 'error': str(e)})

@app.route('/api/memory/projects/<int:project_id>/memory-hierarchy', methods=['GET'])
def get_memory_hierarchy(project_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT tractatus_tree, tractatus_tier FROM memory_projects WHERE id = %s", (project_id,))
        project = cur.fetchone()
        cur.execute("SELECT tier, tree, node_count, created_at FROM tractatus_archive WHERE project_id = %s ORDER BY tier, created_at", (project_id,))
        archives = [{'tier': r[0], 'tree': r[1], 'node_count': r[2], 'created_at': r[3].isoformat()} for r in cur.fetchall()]
        meta_trees = []
        try:
            cur.execute("SELECT id, tree, node_count, archive_start_id, archive_end_id, created_at FROM meta_tractatus WHERE project_id = %s ORDER BY created_at", (project_id,))
            meta_trees = [{'id': r[0], 'tree': r[1], 'node_count': r[2], 'archive_start_id': r[3], 'archive_end_id': r[4], 'created_at': r[5].isoformat()} for r in cur.fetchall()]
        except Exception:
            conn.rollback()
        cur.close()
        conn.close()
        return jsonify({
            'current_tree': project[0] if project else {},
            'current_tier': project[1] if project else 1,
            'archives': archives,
            'meta_trees': meta_trees
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def load_tiered_memory(project_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT tractatus_tree FROM memory_projects WHERE id = %s", (project_id,))
        current = cur.fetchone()
        cur.execute(
            "SELECT tier, tree FROM tractatus_archive WHERE project_id = %s ORDER BY tier DESC",
            (project_id,)
        )
        archives = cur.fetchall()
        meta_trees = []
        try:
            cur.execute("SELECT tree FROM meta_tractatus WHERE project_id = %s ORDER BY created_at DESC LIMIT 3", (project_id,))
            meta_trees = cur.fetchall()
        except Exception:
            conn.rollback()
        cur.close()
        conn.close()

        memory_parts = []
        if meta_trees:
            for i, (mtree,) in enumerate(meta_trees):
                tree_str = json.dumps(mtree, indent=1) if isinstance(mtree, dict) else str(mtree)
                memory_parts.append(f"META-MEMORY (deep synthesis #{i+1} - spans thousands of exchanges):\n{tree_str[:8000]}")
        if current and current[0]:
            tree_str = json.dumps(current[0], indent=1)
            memory_parts.append(f"CURRENT MEMORY (Tier 1 - most recent):\n{tree_str[:12000]}")
        for tier, tree in archives:
            tier_budget = 6000 if tier == 2 else 3000
            tree_str = json.dumps(tree, indent=1)
            memory_parts.append(f"MEMORY TIER {tier} (older, compressed):\n{tree_str[:tier_budget]}")
        return "\n\n".join(memory_parts) if memory_parts else ""
    except Exception as e:
        print(f"Error loading tiered memory: {e}")
        return ""

def load_cross_session_context(project_id, current_session_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(
            "SELECT id, title, transcript FROM memory_sessions WHERE project_id = %s AND id != %s ORDER BY created_at DESC LIMIT 5",
            (project_id, current_session_id)
        )
        sessions = cur.fetchall()
        cur.close()
        conn.close()

        context_parts = []
        total_chars = 0
        for sid, title, transcript in sessions:
            if not transcript:
                continue
            session_text = f"--- Previous session: {title} ---\n"
            for msg in transcript[-6:]:
                role = msg.get('role', 'user')
                content = msg.get('content', '')[:2000]
                session_text += f"{role}: {content}\n"
            if total_chars + len(session_text) > 15000:
                break
            context_parts.append(session_text)
            total_chars += len(session_text)
        return "\n".join(context_parts)
    except Exception as e:
        print(f"Error loading cross-session context: {e}")
        return ""

def update_tractatus_tree(project_id, user_message, assistant_response, thinker):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT tractatus_tree FROM memory_projects WHERE id = %s", (project_id,))
        r = cur.fetchone()
        current_tree = r[0] if r and r[0] else {}
        cur.close()
        conn.close()

        tree_json = json.dumps(current_tree, indent=1)[:8000] if current_tree else "{}"

        analysis_prompt = f"""Analyze this conversation exchange and generate updates for a Tractatus-style knowledge tree.

EXISTING TREE:
{tree_json}

NEW EXCHANGE:
User: {user_message[:3000]}
{thinker}: {assistant_response[:5000]}

Generate a JSON object with Wittgenstein-style numbered keys (like "1.0", "1.1", "1.11", "2.0", etc.).
Each value must be a tagged statement using one of these tags:
- ASSERTS: [confirmed fact or claim]
- REJECTS: [refuted claim]
- ASSUMES: [working assumption]
- OPEN: [unresolved question]
- RESOLVED: [previously open question now answered]

Rules:
1. Build on existing tree numbering if present
2. Add new branches for new topics
3. Mark resolved items if previous OPEN questions were answered
4. Keep statements concise but precise
5. Only include genuinely substantive points from the exchange

Return ONLY valid JSON, no markdown, no explanation."""

        if not anthropic_client:
            return

        response = anthropic_client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=4000,
            messages=[{"role": "user", "content": analysis_prompt}]
        )

        response_text = response.content[0].text.strip()
        if response_text.startswith('```'):
            response_text = response_text.split('\n', 1)[1] if '\n' in response_text else response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]

        try:
            new_nodes = json.loads(response_text)
        except json.JSONDecodeError:
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                new_nodes = json.loads(json_match.group())
            else:
                print("Could not parse tractatus update")
                return

        merged_tree = {**current_tree, **new_nodes}

        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute(
            "UPDATE memory_projects SET tractatus_tree = %s WHERE id = %s",
            (json.dumps(merged_tree), project_id)
        )
        conn.commit()

        node_count = len(merged_tree)
        print(f"📝 Tractatus updated: {node_count} nodes (added {len(new_nodes)} new)")

        if node_count >= 500:
            print(f"🗜️ Tractatus compression triggered at {node_count} nodes")
            compress_tractatus(project_id, cur, conn)

        cur.close()
        conn.close()

    except Exception as e:
        print(f"Error updating tractatus: {e}")
        import traceback
        traceback.print_exc()

def compress_tractatus(project_id, cur=None, conn=None):
    close_conn = False
    try:
        if not conn:
            conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
            cur = conn.cursor()
            close_conn = True

        cur.execute("SELECT tractatus_tree, tractatus_tier FROM memory_projects WHERE id = %s", (project_id,))
        r = cur.fetchone()
        if not r:
            return
        tree, tier = r

        cur.execute(
            "INSERT INTO tractatus_archive (project_id, tier, tree, node_count) VALUES (%s, %s, %s, %s)",
            (project_id, tier, json.dumps(tree), len(tree))
        )

        tree_json = json.dumps(tree, indent=1)
        compress_prompt = f"""Compress this Tractatus knowledge tree from ~{len(tree)} nodes down to ~100 summary nodes.
Preserve the most important assertions, open questions, and key conclusions.
Use the same Wittgenstein-style numbering and tagging format.

CURRENT TREE:
{tree_json[:30000]}

Return ONLY valid JSON with ~100 nodes, no markdown, no explanation."""

        if anthropic_client:
            response = anthropic_client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=8000,
                messages=[{"role": "user", "content": compress_prompt}]
            )
            response_text = response.content[0].text.strip()
            if response_text.startswith('```'):
                response_text = response_text.split('\n', 1)[1] if '\n' in response_text else response_text[3:]
                if response_text.endswith('```'):
                    response_text = response_text[:-3]
            try:
                compressed = json.loads(response_text)
            except json.JSONDecodeError:
                json_match = re.search(r'\{[\s\S]*\}', response_text)
                compressed = json.loads(json_match.group()) if json_match else {}
        else:
            compressed = {}

        cur.execute(
            "UPDATE memory_projects SET tractatus_tree = %s WHERE id = %s",
            (json.dumps(compressed), project_id)
        )
        conn.commit()
        print(f"🗜️ Compressed from {len(tree)} to {len(compressed)} nodes")

        gevent.spawn(check_meta_tractatus, project_id)

        if close_conn:
            cur.close()
            conn.close()

    except Exception as e:
        print(f"Error compressing tractatus: {e}")
        import traceback
        traceback.print_exc()

def check_meta_tractatus(project_id):
    try:
        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()

        cur.execute("""
            CREATE TABLE IF NOT EXISTS meta_tractatus (
                id SERIAL PRIMARY KEY,
                project_id INTEGER REFERENCES memory_projects(id) ON DELETE CASCADE,
                tree JSONB DEFAULT '{}',
                node_count INTEGER DEFAULT 0,
                archive_start_id INTEGER,
                archive_end_id INTEGER,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        conn.commit()

        cur.execute("SELECT COALESCE(MAX(archive_end_id), 0) FROM meta_tractatus WHERE project_id = %s", (project_id,))
        last_covered = cur.fetchone()[0] or 0

        cur.execute(
            "SELECT id, tree, node_count FROM tractatus_archive WHERE project_id = %s AND id > %s ORDER BY id",
            (project_id, last_covered)
        )
        uncovered = cur.fetchall()

        if len(uncovered) >= 10:
            batch = uncovered[:10]
            combined_trees = {}
            for _, atree, _ in batch:
                if isinstance(atree, str):
                    atree = json.loads(atree)
                for k, v in (atree or {}).items():
                    combined_trees[k] = v

            meta_prompt = f"""You are creating a META-TRACTATUS: a higher-order knowledge tree that synthesizes 10 archived Tractatus trees into one unified summary.

These 10 trees each summarize ~500 nodes of conversation, so this meta-tree represents approximately 5000 exchanges worth of accumulated knowledge.

Your goal: Distill the MOST IMPORTANT patterns, conclusions, recurring themes, key assertions, open questions, and intellectual evolution across all 10 trees.

Use Wittgenstein-style numbered nodes (1, 1.1, 1.2, 2, 2.1, etc.).
Tag each node: ASSERTS:, REJECTS:, ASSUMES:, OPEN:, RESOLVED:, or SYNTHESIZES: (new tag for cross-tree patterns).

Target: ~150 nodes that capture the essential intellectual trajectory.

THE 10 ARCHIVED TREES:
{json.dumps(combined_trees, indent=1)[:40000]}

Return ONLY valid JSON with ~150 nodes, no markdown, no explanation."""

            if anthropic_client:
                response = anthropic_client.messages.create(
                    model="claude-sonnet-4-5-20250929",
                    max_tokens=12000,
                    messages=[{"role": "user", "content": meta_prompt}]
                )
                response_text = response.content[0].text.strip()
                if response_text.startswith('```'):
                    response_text = response_text.split('\n', 1)[1] if '\n' in response_text else response_text[3:]
                    if response_text.endswith('```'):
                        response_text = response_text[:-3]
                try:
                    meta_tree = json.loads(response_text)
                except json.JSONDecodeError:
                    json_match = re.search(r'\{[\s\S]*\}', response_text)
                    meta_tree = json.loads(json_match.group()) if json_match else {}
            else:
                meta_tree = {}

            cur.execute(
                "INSERT INTO meta_tractatus (project_id, tree, node_count, archive_start_id, archive_end_id) VALUES (%s, %s, %s, %s, %s)",
                (project_id, json.dumps(meta_tree), len(meta_tree), batch[0][0], batch[-1][0])
            )
            conn.commit()
            print(f"🧬 Meta-Tractatus generated: {len(meta_tree)} nodes spanning archives {batch[0][0]}-{batch[-1][0]}")

        cur.close()
        conn.close()
    except Exception as e:
        print(f"Error generating meta-tractatus: {e}")
        import traceback
        traceback.print_exc()

@app.route('/api/memory/ask', methods=['POST'])
def memory_ask():
    try:
        data = request.json
        question = data.get('question', '')
        provider = data.get('provider', 'deepseek')
        model = data.get('model', '')
        database = data.get('database', 'freud')
        enhanced_mode = data.get('enhanced_mode', False)
        try:
            creativity_level = min(max(int(data.get('creativity_level', 10)), 1), 20)
        except (TypeError, ValueError):
            creativity_level = 10
        gen_temperature = creativity_temperature(creativity_level)
        answer_length = min(max(data.get('answer_length', 250), 100), 8000)
        quote_count = min(max(data.get('quote_count', 5), 1), 50)
        response_mode = data.get('response_mode', 'standard')
        data_source = data.get('data_source', 'combined')
        elevenlabs_mode = bool(data.get('elevenlabs_mode', False))
        project_id = data.get('project_id')
        session_id = data.get('session_id')

        if not question or not project_id or not session_id:
            return jsonify({'error': 'Question, project_id, and session_id required'}), 400

        if response_mode not in ('dialogue', 'conversation'):
            answer_length, quote_count = detect_explicit_requirements(question, answer_length, quote_count)

        if answer_length >= 2000:
            base_top_k = 30
            rag_top_k = 20
        elif answer_length >= 1500:
            base_top_k = 25
            rag_top_k = 15
        elif answer_length >= 1000:
            base_top_k = 20
            rag_top_k = 10
        elif answer_length >= 500:
            base_top_k = 15
            rag_top_k = 8
        else:
            base_top_k = 10
            rag_top_k = 5

        print(f"[MEMORY MODE] Question: {question[:100]}...")
        print(f"Project: {project_id}, Session: {session_id}, Thinker: {database}")

        if database not in databases:
            return jsonify({'error': f'Database "{database}" not available'}), 400

        searcher = databases[database]

        tiered_memory = load_tiered_memory(project_id)
        cross_session = load_cross_session_context(project_id, session_id)

        conn = psycopg2.connect(os.environ.get('DATABASE_URL'))
        cur = conn.cursor()
        cur.execute("SELECT transcript FROM memory_sessions WHERE id = %s", (session_id,))
        r = cur.fetchone()
        session_transcript = r[0] if r and r[0] else []
        cur.close()
        conn.close()

        transcript_text = ""
        for msg in session_transcript[-20:]:
            role = msg.get('role', 'user')
            content = msg.get('content', '')[:3000]
            transcript_text += f"{role}: {content}\n"

        retrieval_metadata = {}
        try:
            if data_source == 'classic':
                search_result = searcher.search(question, top_k=base_top_k, return_metadata=True)
                relevant_positions = search_result['results']
                retrieval_metadata = search_result['metadata']
                print(f"📚 [MEMORY] Using CLASSIC search - top_k={base_top_k}")
            elif data_source == 'newdb':
                relevant_positions = search_positions_postgres(question, database, top_k=base_top_k)
                quote_results = search_quotes_postgres(question, database, top_k=rag_top_k)
                relevant_positions = merge_position_results(relevant_positions, quote_results, max_results=base_top_k)
                print(f"📊 [MEMORY] Using NEW DB search - top_k={base_top_k}")
            else:
                search_result = searcher.search(question, top_k=base_top_k, return_metadata=True)
                tier1_results = search_result['results']
                retrieval_metadata = search_result['metadata']
                tier2_results = search_positions_postgres(question, database, top_k=base_top_k)
                relevant_positions = merge_position_results(tier1_results, tier2_results, max_results=base_top_k + 5)
                print(f"🔀 [MEMORY] Using COMBINED search - top_k={base_top_k}")
            if retrieval_metadata:
                print(f"   📊 SCANNED {retrieval_metadata.get('total_positions_scanned', 'N/A')} positions, {retrieval_metadata.get('positions_above_threshold', 'N/A')} above threshold")
        except Exception as e:
            print(f"Search error: {e}")
            import traceback
            traceback.print_exc()
            relevant_positions = []

        if hasattr(searcher, 'expand_with_context') and relevant_positions:
            try:
                original_count = len(relevant_positions)
                relevant_positions = searcher.expand_with_context(relevant_positions, max_context=3)
                context_added = len(relevant_positions) - original_count
                if context_added > 0:
                    retrieval_metadata['context_positions_added'] = context_added
                    retrieval_metadata['total_positions_used'] = len(relevant_positions)
                    print(f"   📖 [MEMORY] Expanded {original_count} → {len(relevant_positions)} with argument context")
            except Exception as e:
                print(f"Context expansion error: {e}")

        rag_chunks = []
        try:
            db_url = os.environ.get('DATABASE_URL')
            if db_url:
                keywords = [w for w in question.lower().split() if len(w) > 3][:5]
                rag_conn = psycopg2.connect(db_url)
                rag_cur = rag_conn.cursor()
                search_query = ' | '.join(keywords)
                rag_cur.execute("""
                    SELECT source_file, chunk_text, ts_rank(to_tsvector('english', chunk_text), to_tsquery('english', %s)) as rank
                    FROM text_chunks WHERE thinker = %s AND to_tsvector('english', chunk_text) @@ to_tsquery('english', %s)
                    ORDER BY rank DESC LIMIT %s
                """, (search_query, database, search_query, rag_top_k))
                rag_chunks = [{'source_file': r[0], 'chunk_text': r[1], 'rank': r[2]} for r in rag_cur.fetchall()]
                rag_cur.close()
                rag_conn.close()
        except Exception as e:
            print(f"RAG search error: {e}")

        low_relevance = not relevant_positions or (retrieval_metadata.get('max_similarity', 0) < 0.25)

        deduced_rules = []
        thinker_name = {
            'freud': 'Sigmund Freud', 'kuczynski': 'Zbigniew Hippolyte Iwo Kuczynski',
            'jung': 'Carl Gustav Jung', 'hume': 'David Hume',
            'nietzsche': 'Friedrich Nietzsche', 'bergler': 'Edmund Bergler'
        }.get(database, database.capitalize())

        prompt = build_prompt(question, relevant_positions, database, transcript_text, enhanced_mode, low_relevance, deduced_rules, answer_length, quote_count, response_mode, rag_chunks, creativity_level)

        memory_section = ""
        if tiered_memory:
            memory_section = f"\n\nPROJECT MEMORY (Tractatus Tree - what you remember from past conversations):\n{tiered_memory}\n"
            print(f"📝 [MEMORY] Tiered memory loaded: {len(tiered_memory)} chars")
        else:
            print(f"📝 [MEMORY] No tiered memory found")
        if cross_session:
            memory_section += f"\nPREVIOUS CONVERSATIONS IN THIS PROJECT:\n{cross_session}\n"
            print(f"📝 [MEMORY] Cross-session context loaded: {len(cross_session)} chars")
        else:
            print(f"📝 [MEMORY] No cross-session context found")

        if memory_section:
            if "Begin your response" in prompt:
                prompt = prompt.replace("Begin your response", f"{memory_section}\nYou REMEMBER previous conversations in this project. Use that memory to maintain continuity. If the user asks what you discussed before, TELL THEM based on the previous conversations above.\n\nBegin your response", 1)
                print(f"✅ [MEMORY] Memory section injected into prompt ({len(memory_section)} chars)")
            else:
                prompt += f"\n{memory_section}\n"
                print(f"⚠️ [MEMORY] 'Begin your response' not found in prompt, appended memory to end")

        accumulated_text = [""]
        full_answer = []
        last_saved_word_count = [0]

        def count_words_and_pause(text):
            accumulated_text[0] += text
            return []

        def try_provider(prov, prompt_text, sys_prompt):
            if prov == 'anthropic':
                if not anthropic_client:
                    raise Exception("Anthropic not configured")
                model_name = model if model and prov == provider else "claude-sonnet-4-5-20250929"
                if model_name not in ("claude-sonnet-4-5-20250929", "claude-opus-4-5"):
                    model_name = "claude-sonnet-4-5-20250929"
                with anthropic_client.messages.stream(
                    model=model_name, max_tokens=16000, temperature=gen_temperature, system=sys_prompt,
                    messages=[{"role": "user", "content": prompt_text}]
                ) as stream:
                    for text in stream.text_stream:
                        full_answer.append(text)
                        yield f"data: {json.dumps({'type': 'token', 'data': text})}\n\n"
                        count_words_and_pause(text)
            elif prov == 'openai':
                if not openai_client:
                    raise Exception("OpenAI not configured")
                model_name = model if model and prov == provider else "gpt-4o"
                stream = openai_client.chat.completions.create(
                    model=model_name, messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": prompt_text}],
                    stream=True, max_tokens=16000, temperature=gen_temperature, timeout=180.0
                )
                for chunk in stream:
                    if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_answer.append(content)
                        yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                        count_words_and_pause(content)
            elif prov == 'deepseek':
                if not deepseek_client:
                    raise Exception("DeepSeek not configured")
                model_name = model if model and prov == provider else "deepseek-chat"
                stream = deepseek_client.chat.completions.create(
                    model=model_name, messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": prompt_text}],
                    stream=True, max_tokens=16000, temperature=gen_temperature, timeout=300.0
                )
                for chunk in stream:
                    if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_answer.append(content)
                        yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                        count_words_and_pause(content)
            elif prov == 'venice':
                if not venice_client:
                    raise Exception("Venice not configured")
                model_name = model if model and prov == provider else "llama-3.3-70b"
                stream = venice_client.chat.completions.create(
                    model=model_name, messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": prompt_text}],
                    stream=True, max_tokens=16000, temperature=gen_temperature, timeout=300.0
                )
                for chunk in stream:
                    if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_answer.append(content)
                        yield f"data: {json.dumps({'type': 'token', 'data': content})}\n\n"
                        count_words_and_pause(content)

        if response_mode == 'dialogue':
            sys_prompt = f"You are {thinker_name} in dialogue. Brief but substantive responses drawn from your writings. NEVER refuse to answer."
        else:
            sys_prompt = f"MANDATORY WORD COUNT: {answer_length} WORDS MINIMUM. You are {thinker_name}. Write plain prose, no markdown, no source citations. Do NOT stop before {answer_length} words."
        if elevenlabs_mode and response_mode in ('dialogue', 'conversation'):
            sys_prompt += ELEVENLABS_DIRECTIVE

        def generate():
            event_queue = gevent_queue.Queue()
            streaming_done = [False]
            HEARTBEAT_INTERVAL = 3

            def heartbeat_worker():
                hb_count = 0
                while not streaming_done[0]:
                    gevent.sleep(HEARTBEAT_INTERVAL)
                    if not streaming_done[0]:
                        hb_count += 1
                        event_queue.put(f": keepalive {hb_count}\n\n")

            def content_worker():
                try:
                    sources = [p['position_id'] for p in relevant_positions]
                    positions_data = [{
                        'id': p.get('position_id', ''),
                        'text': (p.get('text') or p.get('thesis', ''))[:300],
                        'domain': p.get('domain', 'philosophy'),
                        'similarity': p.get('similarity', 0)
                    } for p in relevant_positions]
                    event_queue.put(f"data: {json.dumps({'type': 'sources', 'data': sources, 'positions': positions_data})}\n\n")

                    if retrieval_metadata:
                        event_queue.put(f"data: {json.dumps({'type': 'retrieval_log', 'data': retrieval_metadata})}\n\n")

                    providers_to_try = [provider] + get_fallback_providers(provider)
                    success = False

                    for prov in providers_to_try:
                        if success:
                            break
                        try:
                            for event in try_provider(prov, prompt, sys_prompt):
                                event_queue.put(event)
                            success = True
                            print(f"✓ [MEMORY] {prov} completed: {len(accumulated_text[0].split())} words")
                        except Exception as e:
                            print(f"✗ [MEMORY] {prov} failed: {e}")
                            continue

                    if not success:
                        event_queue.put(f"data: {json.dumps({'type': 'token', 'data': 'All AI providers unavailable. Please try again.'})}\n\n")

                    answer_text = ''.join(full_answer)
                    if answer_text:
                        try:
                            conn2 = psycopg2.connect(os.environ.get('DATABASE_URL'))
                            cur2 = conn2.cursor()
                            cur2.execute("SELECT transcript FROM memory_sessions WHERE id = %s", (session_id,))
                            r2 = cur2.fetchone()
                            transcript = r2[0] if r2 and r2[0] else []
                            transcript.append({'role': 'user', 'content': question})
                            transcript.append({'role': 'assistant', 'content': answer_text})
                            cur2.execute("UPDATE memory_sessions SET transcript = %s WHERE id = %s", (json.dumps(transcript), session_id))
                            conn2.commit()
                            cur2.close()
                            conn2.close()
                            print(f"💾 [MEMORY] Transcript saved ({len(transcript)} messages)")
                        except Exception as e:
                            print(f"Error saving transcript: {e}")

                        gevent.spawn(update_tractatus_tree, project_id, question, answer_text, thinker_name)

                    event_queue.put(f"data: {json.dumps({'type': 'done'})}\n\n")
                except Exception as e:
                    print(f"[MEMORY] Content worker error: {e}")
                    import traceback
                    traceback.print_exc()
                    event_queue.put(f"data: {json.dumps({'type': 'done'})}\n\n")
                finally:
                    streaming_done[0] = True
                    event_queue.put(None)

            heartbeat_greenlet = gevent.spawn(heartbeat_worker)
            content_greenlet = gevent.spawn(content_worker)

            try:
                while True:
                    event = event_queue.get()
                    if event is None:
                        break
                    yield event
            finally:
                streaming_done[0] = True
                heartbeat_greenlet.kill()
                content_greenlet.kill()

        return Response(
            generate(),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache, no-transform',
                'X-Accel-Buffering': 'no',
                'Connection': 'keep-alive'
            }
        )
    except Exception as e:
        print(f"ERROR in /api/memory/ask: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    total_positions = 24382
    print("\n" + "="*60)
    print("  FreudGPT - Multi-Thinker AI Assistant (Lazy Loading)")
    print("="*60)
    print(f"  Available databases: {', '.join(databases.keys())}")
    print(f"  Total philosophical positions: {total_positions}")
    print(f"  Data loads on first query (memory-optimized startup)")
    print(f"  Server starting on http://0.0.0.0:{port}")
    print("="*60 + "\n")
    app.run(host='0.0.0.0', port=port, debug=False)